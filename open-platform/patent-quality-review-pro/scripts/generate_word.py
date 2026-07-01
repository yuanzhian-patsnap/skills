"""
专利申请文件质量评价表 Word 生成脚本 v4（V2格式）
格式规格（依据 专利申请文件质量评价表V2.docx）：
  页面：A4 (21×29.7cm)，四边距均 2.0cm
  字体：宋体，全文统一（含 w:eastAsia 属性，确保中文正确渲染）
  字色：全文纯黑色 (RGB 0,0,0)，禁止任何彩色字体
  字号：正文 10.5pt，大标题 14pt 加粗居中，表标题 10.5pt 加粗居中
  表底色：全表无底色（禁止灰色/彩色底色）
  表格边框：黑色单实线

输出结构（2页）：
  第1页：主表（基本信息 + 评分情况，整合为一张表）
  第2页：附表1 主要问题与修改建议

主表结构：
  行1: "基本信息"（跨列标题，加粗居中）
  行2: 专利名称
  行3: 技术领域
  行4: 代理机构名称
  行5: 联系人（代理机构）| 联系电话
  行6: "评分情况"（跨列标题，加粗居中）
  行7: 一、权利要求书质量（子标题行）
  行8: 评价指标 | 得分 | 加权得分 | 主要问题描述（列标题行）
  行9: 技术方案表述（10%）
  行10: 保护范围规划（10%）
  行11: 权利稳定性（10%）
  行12: 权利要求布局（20%）
  行13: 二、说明书质量（子标题行）
  行14: 评价指标 | 得分 | 加权得分 | 主要问题描述（列标题行）
  行15: 背景技术论述（10%）
  行16: 技术问题与技术效果论述（10%）
  行17: 说明书充分公开情况（20%）
  行18: 三、其他（子标题行）
  行19: 评价指标 | 得分 | 加权得分 | 主要问题描述（列标题行）
  行20: 形式问题（10%）
  行21: 评价总分（跨列，居中加粗）

附表1：序号 | 严重程度 | 问题名称 | 涉及位置 | 问题描述 | 修改建议
  最后行：评价人姓名 | 评价日期 | 签字/盖章

数据读取方式（优先级从高到低）：
  1. 命令行参数 --data <json_file_path>  ← 推荐，最稳定
  2. 环境变量 REVIEW_DATA_FILE 指向 JSON 文件路径
  3. 环境变量 REVIEW_DATA_JSON 内嵌 JSON 字符串（备用）
  4. 以上均无则生成空白模板

输出路径（优先级从高到低）：
  1. 命令行参数 --output <file_path>
  2. 环境变量 EUREKA_PYTHON_OUTPUT_WORD_DOC
  3. 脚本同级目录 ../output/专利申请文件质量评价表+未命名.docx

scores 数据结构（8项四元组）：
  scores[0] = ("技术方案表述（10%）",           原始分, 加权分, 问题描述)
  scores[1] = ("保护范围规划（10%）",            原始分, 加权分, 问题描述)
  scores[2] = ("权利稳定性（10%）",              原始分, 加权分, 问题描述)
  scores[3] = ("权利要求布局（20%）",            原始分, 加权分, 问题描述)
  scores[4] = ("背景技术论述（10%）",            原始分, 加权分, 问题描述)
  scores[5] = ("技术问题与技术效果论述（10%）",  原始分, 加权分, 问题描述)
  scores[6] = ("说明书充分公开情况（20%）",      原始分, 加权分, 问题描述)
  scores[7] = ("形式问题（10%）",               原始分, 加权分, 问题描述)

  注：V2格式将权利要求书质量拆为4项（技术方案表述/保护范围规划/权利稳定性/权利要求布局），
      说明书质量3项，其他1项，共8项。

issues 数据结构（每项六元组）：
  (序号, 严重程度, 问题名称, 涉及位置, 问题描述, 修改建议)

Usage:
  python generate_word.py --data /path/to/review_data.json --output /path/to/output.docx
"""
import os
import sys
import json
import argparse
import re
import subprocess
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 参数解析 ──────────────────────────────────────────────────────────────────
parser = argparse.ArgumentParser(description="生成专利申请文件质量评价表 Word 文件（V2格式）")
parser.add_argument("--data",   help="评审数据 JSON 文件路径")
parser.add_argument("--output", help="输出 Word 文件路径（含文件名）")
parser.add_argument("--no-open-output", action="store_true", help="生成后不自动打开输出目录（测试用）")
args, _ = parser.parse_known_args()

# ── 输出路径（阶段1：仅处理显式参数，兜底延迟到数据读取后）────────────────────
OUTPUT_PATH = (
    args.output
    or os.environ.get("EUREKA_PYTHON_OUTPUT_WORD_DOC", "")
)
# OUTPUT_PATH 的兜底（阶段2）在数据读取完成后执行，见下方 "输出路径阶段2"

FONT_NAME = "宋体"
WORD_JOINER = "\u2060"

# ── 数据读取 ──────────────────────────────────────────────────────────────────
DATA = None
DATA_WAS_PROVIDED = False
DATA_FILE_PATH = ""

if args.data and os.path.exists(args.data):
    DATA_FILE_PATH = os.path.abspath(args.data)
    with open(args.data, "r", encoding="utf-8") as f:
        DATA = json.load(f)
    DATA_WAS_PROVIDED = True

if DATA is None:
    data_file = os.environ.get("REVIEW_DATA_FILE", "")
    if data_file and os.path.exists(data_file):
        DATA_FILE_PATH = os.path.abspath(data_file)
        with open(data_file, "r", encoding="utf-8") as f:
            DATA = json.load(f)
        DATA_WAS_PROVIDED = True

if DATA is None:
    raw = os.environ.get("REVIEW_DATA_JSON", "")
    if raw:
        try:
            DATA = json.loads(raw)
            DATA_WAS_PROVIDED = True
        except Exception:
            DATA = None

# 兜底：空白模板（V2格式，8项指标）
if DATA is None:
    DATA = {
        "basic": {
            "patent_name": "（专利名称）",
            "tech_field":  "（技术领域）",
            "agency":      "（代理机构名称）",
            "agent":       "（联系人）",
            "agent_tel":   "—",
        },
        "scores": [
            ("技术方案表述（10%）",           "", "", ""),
            ("保护范围规划（10%）",            "", "", ""),
            ("权利稳定性（10%）",              "", "", ""),
            ("权利要求布局（20%）",            "", "", ""),
            ("背景技术论述（10%）",            "", "", ""),
            ("技术问题与技术效果论述（10%）",  "", "", ""),
            ("说明书充分公开情况（20%）",      "", "", ""),
            ("形式问题（10%）",               "", "", ""),
        ],
        "total_score": "",
        "conclusion":  "",
        "issues":      [],
        "review_date": "____年____月____日",
        "reviewer_name": "AI辅助评审（Eureka）",
    }

# ── 输出路径阶段2：若未通过参数/环境变量指定，则用专利名称拼接兜底路径 ──────
if not OUTPUT_PATH:
    _patent_name = ""
    if DATA and isinstance(DATA.get("basic"), dict):
        _patent_name = str(DATA["basic"].get("patent_name", "")).strip()
        # 去掉占位符
        if _patent_name in {"（专利名称）", "(专利名称)", ""}:
            _patent_name = ""
    _safe_name = _patent_name if _patent_name else "未命名"
    OUTPUT_PATH = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "..", "output",
        f"专利申请文件质量评价表+{_safe_name}.docx"
    )
if not OUTPUT_PATH.lower().endswith(".docx"):
    OUTPUT_PATH = f"{OUTPUT_PATH}.docx"
os.makedirs(os.path.dirname(os.path.abspath(OUTPUT_PATH)), exist_ok=True)


# ── 生成前证据链硬校验 ────────────────────────────────────────────────────────

NOVELTY_ARTIFACTS = [
    "claim_elements.md",
    "prior_art_catalog.json",
    "element_mapping.md",
    "claim_diff_matrix.md",
    "novelty_report.md",
]

NON_OBVIOUSNESS_ARTIFACTS = [
    "jurisdiction_plan.md",
    "claim_diff_matrix.md",
    "combination_candidates.json",
    "motivation_matrix.md",
    "secondary_considerations.md",
    "inventive_step_report.md",
]

NON_OBVIOUSNESS_D2D3_ARTIFACTS = [
    "combination_candidates.json",
    "motivation_matrix.md",
    "inventive_step_report.md",
]

UNREAD_REFERENCE_MARKERS = [
    "摘要读取",
    "仅读取摘要",
    "未fetch",
    "未 fetch",
    "未read",
    "未 read",
    "未读取全文",
    "未fetch/read全文",
]


def infer_session_root():
    """Infer the Eureka session folder from data/output paths or cwd."""
    seeds = [DATA_FILE_PATH, os.path.abspath(OUTPUT_PATH), os.getcwd()]
    for seed in seeds:
        if not seed:
            continue
        path = os.path.abspath(seed)
        if os.path.isfile(path):
            path = os.path.dirname(path)
        while True:
            if os.path.basename(path).startswith("session-"):
                return path
            if os.path.exists(os.path.join(path, "session.runtime.json")):
                return path
            parent = os.path.dirname(path)
            if parent == path:
                break
            path = parent
    return os.getcwd()


SESSION_ROOT = infer_session_root()


def resolve_artifact_path(raw_path):
    """Resolve Eureka-style @session/@skill_workspace refs and plain paths."""
    if not raw_path or not isinstance(raw_path, str):
        return ""
    raw_path = raw_path.strip()
    if raw_path.startswith("@session/"):
        return os.path.join(SESSION_ROOT, raw_path[len("@session/"):])
    if raw_path == "@session":
        return SESSION_ROOT
    if raw_path.startswith("@skill_workspace/"):
        return os.path.join(os.getcwd(), raw_path[len("@skill_workspace/"):])
    if raw_path == "@skill_workspace":
        return os.getcwd()
    if os.path.isabs(raw_path):
        return raw_path

    cwd_candidate = os.path.abspath(raw_path)
    if os.path.exists(cwd_candidate):
        return cwd_candidate
    return os.path.join(SESSION_ROOT, raw_path)


def artifact_candidates(name, category=None):
    paths = []
    artifact_paths = DATA.get("artifact_paths", {})
    stem = os.path.splitext(name)[0]

    def add(value):
        if isinstance(value, str):
            paths.append(value)
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, str):
                    paths.append(item)

    if isinstance(artifact_paths, dict):
        if category and isinstance(artifact_paths.get(category), dict):
            add(artifact_paths[category].get(name))
            add(artifact_paths[category].get(stem))
        add(artifact_paths.get(name))
        add(artifact_paths.get(stem))

        # non-obviousness-check may reuse the novelty claim_diff_matrix.
        if name == "claim_diff_matrix.md" and category == "non_obviousness":
            novelty_paths = artifact_paths.get("novelty")
            if isinstance(novelty_paths, dict):
                add(novelty_paths.get(name))
                add(novelty_paths.get(stem))
    elif isinstance(artifact_paths, list):
        add(artifact_paths)

    for item in DATA.get("evidence_artifacts", []):
        if isinstance(item, str) and item.strip().endswith(name):
            add(item)

    return paths


def existing_nonempty_path(paths):
    for raw_path in paths:
        path = resolve_artifact_path(raw_path)
        if path and os.path.exists(path) and os.path.getsize(path) > 0:
            return path
    return ""


def reference_entries():
    entries = []
    raw_refs = DATA.get("fetched_references", [])
    if isinstance(raw_refs, list):
        for item in raw_refs:
            if isinstance(item, dict):
                entries.append(item)
    raw_paths = DATA.get("fetched_reference_paths", {})
    if isinstance(raw_paths, dict):
        for role, value in raw_paths.items():
            values = value if isinstance(value, list) else [value]
            for path in values:
                if isinstance(path, str):
                    entries.append({"role": role, "path": path})
    elif isinstance(raw_paths, list):
        for path in raw_paths:
            if isinstance(path, str):
                entries.append({"role": "unknown", "path": path})
    return entries


def has_existing_reference(role_names):
    return bool(existing_reference_paths(role_names))


def existing_reference_paths(role_names):
    wanted = {role.upper() for role in role_names}
    paths = []
    for entry in reference_entries():
        role = str(entry.get("role", "")).upper()
        path = entry.get("path", "")
        if role in wanted:
            resolved = resolve_artifact_path(path)
            if resolved and os.path.exists(resolved) and os.path.getsize(resolved) > 0:
                paths.append(resolved)
    return paths


def artifact_has_unread_reference_marker(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
    except Exception:
        return False
    return any(marker in text for marker in UNREAD_REFERENCE_MARKERS)


def session_file_path(filename):
    path = os.path.join(SESSION_ROOT, filename)
    return path if os.path.exists(path) else ""


def actual_loaded_skills():
    """Read Eureka runtime records instead of trusting the review JSON alone."""
    skills = set()

    runtime_path = session_file_path("session.runtime.json")
    if runtime_path:
        try:
            with open(runtime_path, "r", encoding="utf-8") as f:
                runtime = json.load(f)
            skills.update(runtime.get("loaded_skills") or [])
            active_skill = runtime.get("active_skill")
            if active_skill:
                skills.add(active_skill)
        except Exception:
            pass

    index_path = session_file_path("turn_context_index.json")
    if index_path:
        try:
            with open(index_path, "r", encoding="utf-8") as f:
                index = json.load(f)
            for record in index.get("records", []):
                if record.get("tool_name") != "skills":
                    continue
                tool_input = record.get("input") or {}
                result = record.get("result_digest") or {}
                if tool_input.get("action") != "load":
                    continue
                if record.get("status") != "completed":
                    continue
                if result.get("success") is False:
                    continue
                skill_name = tool_input.get("name") or result.get("skill")
                if skill_name:
                    skills.add(skill_name)
        except Exception:
            pass

    return skills


def has_session_runtime_records():
    return bool(session_file_path("session.runtime.json") or session_file_path("turn_context_index.json"))


def validation_error_path():
    return os.path.join(
        os.path.dirname(os.path.abspath(OUTPUT_PATH)),
        "review_validation_errors.json",
    )


def write_validation_errors(errors):
    error_path = validation_error_path()
    with open(error_path, "w", encoding="utf-8") as f:
        json.dump({"errors": errors}, f, ensure_ascii=False, indent=2)
    return error_path


def clear_validation_errors():
    error_path = os.path.abspath(validation_error_path())
    output_dir = os.path.dirname(os.path.abspath(OUTPUT_PATH))
    if os.path.basename(error_path) != "review_validation_errors.json":
        return
    if os.path.dirname(error_path) != output_dir:
        return
    if os.path.exists(error_path):
        os.remove(error_path)


def validate_evidence_contract():
    errors = []
    evidence_mode = DATA.get("evidence_mode")
    novelty = DATA.get("novelty_conclusion")
    inventive = DATA.get("inventive_step_conclusion")
    loaded_skills = set(DATA.get("loaded_skills") or [])
    runtime_skills = actual_loaded_skills()
    enforce_runtime_skills = has_session_runtime_records()

    def require_loaded_skill(skill_name, message):
        if skill_name not in loaded_skills:
            errors.append(message)
        elif enforce_runtime_skills and skill_name not in runtime_skills:
            errors.append(f"{message} session.runtime.json/turn_context_index.json 中未记录实际加载 {skill_name}。")

    if not evidence_mode:
        errors.append("缺少 evidence_mode，无法判断是否完成官方技能证据链。")
    if evidence_mode == "degraded":
        errors.append("evidence_mode=degraded 不允许生成最终 Word；请补齐证据链或改为经用户确认的 document_only。")

    if evidence_mode == "document_only":
        if DATA.get("document_only_confirmed") is not True:
            errors.append("document_only 必须设置 document_only_confirmed=true，表示用户明确确认跳过外部检索/官方技能。")
        if novelty != "uncertain":
            errors.append("document_only 下 novelty_conclusion 只能为 uncertain。")
        if inventive != "uncertain":
            errors.append("document_only 下 inventive_step_conclusion 只能为 uncertain。")
        return errors

    if evidence_mode == "full":
        require_loaded_skill("novelty-check", "evidence_mode=full 时 loaded_skills 必须包含 novelty-check。")
        if novelty not in {"novelty_rejected", "novelty_preserved", "uncertain"}:
            errors.append("novelty_conclusion 必须为 novelty_rejected/novelty_preserved/uncertain。")
        if not has_existing_reference(["D1"]):
            errors.append("evidence_mode=full 时 fetched_references 必须包含一个已 fetch/read 的 D1 原文路径。")

        for filename in NOVELTY_ARTIFACTS:
            if not existing_nonempty_path(artifact_candidates(filename, "novelty")):
                errors.append(f"缺少 novelty-check 官方产物或文件为空：{filename}")

        if novelty != "novelty_rejected":
            require_loaded_skill("non-obviousness-check", "新颖性未被否定时，loaded_skills 必须包含 non-obviousness-check。")
            if inventive not in {"strong", "weak", "uncertain"}:
                errors.append("inventive_step_conclusion 必须为 strong/weak/uncertain。")
            for filename in NON_OBVIOUSNESS_ARTIFACTS:
                if not existing_nonempty_path(artifact_candidates(filename, "non_obviousness")):
                    errors.append(f"缺少 non-obviousness-check 官方产物或文件为空：{filename}")
            if inventive in {"strong", "weak"}:
                d2d3_paths = existing_reference_paths(["D2", "D3"])
                if not d2d3_paths:
                    errors.append("inventive_step_conclusion 为 strong/weak 时，fetched_references 必须包含至少一个已 fetch/read 的 D2 或 D3 原文路径。")
                else:
                    newest_reference_mtime = max(os.path.getmtime(path) for path in d2d3_paths)
                    for filename in NON_OBVIOUSNESS_D2D3_ARTIFACTS:
                        artifact_path = existing_nonempty_path(artifact_candidates(filename, "non_obviousness"))
                        if not artifact_path:
                            continue
                        if os.path.getmtime(artifact_path) + 1 < newest_reference_mtime:
                            errors.append(f"{filename} 早于 D2/D3 原文读取时间，需在 fetch/read D2/D3 后重写 non-obviousness-check 产物。")
                        if artifact_has_unread_reference_marker(artifact_path):
                            errors.append(f"{filename} 仍包含 D2/D3 未阅读全文的表述，不能支撑 strong/weak 结论。")

    return errors


if DATA_WAS_PROVIDED:
    import review_validation as evidence_validation

    validation_errors, _, _ = evidence_validation.validate_evidence_contract(
        DATA,
        data_file_path=DATA_FILE_PATH,
        output_path=OUTPUT_PATH,
    )
    if validation_errors:
        validation_error_path = evidence_validation.write_validation_errors(validation_errors, OUTPUT_PATH)
        print("VALIDATION_ERROR: evidence chain incomplete; Word generation blocked.")
        print(f"Validation details: {validation_error_path}")
        for error in validation_errors:
            print(f"- {error}")
        sys.exit(2)
    evidence_validation.clear_validation_errors(OUTPUT_PATH)

# ── 辅助函数 ──────────────────────────────────────────────────────────────────

def set_page(doc):
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.0)

def set_col_widths(table, widths_cm):
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w_cm in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w_cm * 567)))
        tblGrid.append(gc)

def set_cell_width(cell, width_cm):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcW"))
    if existing is not None:
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def set_borders(table, sz=4):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_run_font(run, font_name=FONT_NAME, font_size_pt=None, bold=False):
    """设置字体：宋体 + 纯黑色（禁止彩色）"""
    run.font.name = font_name
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),    font_name)
    rFonts.set(qn("w:hAnsi"),    font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:hint"),     "eastAsia")
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    run.bold = bold

def add_para(container, text, align=WD_ALIGN_PARAGRAPH.LEFT,
             font_size_pt=None, bold=False):
    para = container.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run(text)
    set_run_font(run, font_size_pt=font_size_pt, bold=bold)
    return para

def normalize_cell_text(text):
    """Normalize report text before Word rendering.

    Chinese Word tables wrap narrow columns aggressively. Half-width and curly
    quotes often become orphan punctuation at the beginning of a wrapped line.
    Use Chinese corner quotes and invisible word joiners to keep the quote with
    the quoted term.
    """
    if text is None:
        return ""
    text = str(text)
    text = re.sub(r"\s*[\r\n]+\s*", " ", text)
    text = re.sub(r'"([^"\n]{1,80})"', lambda m: f"「{m.group(1)}」", text)
    text = re.sub(r"'([^'\n]{1,80})'", lambda m: f"「{m.group(1)}」", text)
    text = (
        text.replace("“", "「")
            .replace("”", "」")
            .replace("‘", "「")
            .replace("’", "」")
            .replace("『", "「")
            .replace("』", "」")
            .replace("＂", "「")
            .replace("＇", "」")
    )
    text = re.sub(
        r"「\s*([^」\n]{1,80}?)\s*」",
        lambda m: f"「{WORD_JOINER}{m.group(1)}{WORD_JOINER}」",
        text,
    )
    text = re.sub(r'(?<=[\u4e00-\u9fff])["\']|["\'](?=[\u4e00-\u9fff])', "", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()

def write_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT,
               font_size_pt=None, bold=False,
               valign=WD_ALIGN_VERTICAL.CENTER):
    text = normalize_cell_text(text)
    cell.vertical_alignment = valign
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run(text)
    set_run_font(run, font_size_pt=font_size_pt, bold=bold)

def add_page_break(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run  = para.add_run()
    br   = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

# ── 构建文档 ──────────────────────────────────────────────────────────────────
doc = Document()
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)
set_page(doc)

B      = DATA["basic"]
SCORES = DATA["scores"]
TOTAL  = str(DATA.get("total_score", ""))
CONCL  = DATA.get("conclusion", "")
ISSUES = DATA.get("issues", [])
RDATE  = DATA.get("review_date", "____年____月____日")
RNAME  = DATA.get("reviewer_name", "AI辅助评审（Eureka）")

# ═══════════════════════════════════════════════════════════
# 第1页：主表（基本信息 + 评分情况，整合为一张表）
# ═══════════════════════════════════════════════════════════

# 文档大标题
add_para(doc, "专利申请文件质量评价表",
         align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=14, bold=True)

# 主表列宽：评价指标列宽，得分，加权得分，主要问题描述
# 4列：指标名 5.0cm | 得分 2.5cm | 加权得分 2.5cm | 主要问题描述 7.0cm  = 17.0cm（A4正文宽）
TM = [5.0, 2.5, 2.5, 7.0]
NC = 4
tbl = doc.add_table(rows=0, cols=NC)
set_col_widths(tbl, TM)
set_borders(tbl)

# ── 基本信息 ──
r = tbl.add_row()
for i, c in enumerate(r.cells): set_cell_width(c, TM[i])
cell = r.cells[0].merge(r.cells[NC - 1])
write_cell(cell, "基本信息",
           align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5, bold=True)

def basic_row(tbl, label, value):
    row = tbl.add_row()
    for i, c in enumerate(row.cells): set_cell_width(c, TM[i])
    write_cell(row.cells[0], label, font_size_pt=10.5)
    m = row.cells[1].merge(row.cells[NC - 1])
    write_cell(m, value, font_size_pt=10.5)
    return row

basic_row(tbl, "专利名称",     B["patent_name"])
basic_row(tbl, "技术领域",     B["tech_field"])
basic_row(tbl, "代理机构名称", B["agency"])

# 联系人 | 联系电话（两列合并）
row = tbl.add_row()
for i, c in enumerate(row.cells): set_cell_width(c, TM[i])
write_cell(row.cells[0], "联系人", font_size_pt=10.5)
write_cell(row.cells[1], B["agent"], font_size_pt=10.5)
write_cell(row.cells[2], "联系电话", font_size_pt=10.5)
write_cell(row.cells[3], B["agent_tel"], font_size_pt=10.5)

# ── 评分情况 ──
r2 = tbl.add_row()
for i, c in enumerate(r2.cells): set_cell_width(c, TM[i])
cell2 = r2.cells[0].merge(r2.cells[NC - 1])
write_cell(cell2, "评分情况",
           align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5, bold=True)

# 列标题行（共用）
def add_col_header(tbl):
    hr = tbl.add_row()
    for i, c in enumerate(hr.cells): set_cell_width(c, TM[i])
    for i, txt in enumerate(["评价指标", "得分", "加权得分", "主要问题描述"]):
        write_cell(hr.cells[i], txt,
                   align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5, bold=True)

# 维度一：权利要求书质量（scores[0-3]）
sr = tbl.add_row()
for i, c in enumerate(sr.cells): set_cell_width(c, TM[i])
sc = sr.cells[0].merge(sr.cells[NC - 1])
write_cell(sc, "一、权利要求书质量", font_size_pt=10.5, bold=False)

add_col_header(tbl)

for idx in range(4):
    name, score, wscore, desc = SCORES[idx]
    row = tbl.add_row()
    for i, c in enumerate(row.cells): set_cell_width(c, TM[i])
    write_cell(row.cells[0], name,        font_size_pt=10.5)
    write_cell(row.cells[1], str(score),  align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5)
    write_cell(row.cells[2], str(wscore), align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5)
    write_cell(row.cells[3], str(desc),   font_size_pt=10.5)

# 维度二：说明书质量（scores[4-6]）
sr2 = tbl.add_row()
for i, c in enumerate(sr2.cells): set_cell_width(c, TM[i])
sc2 = sr2.cells[0].merge(sr2.cells[NC - 1])
write_cell(sc2, "二、说明书质量", font_size_pt=10.5, bold=False)

add_col_header(tbl)

for idx in range(4, 7):
    name, score, wscore, desc = SCORES[idx]
    row = tbl.add_row()
    for i, c in enumerate(row.cells): set_cell_width(c, TM[i])
    write_cell(row.cells[0], name,        font_size_pt=10.5)
    write_cell(row.cells[1], str(score),  align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5)
    write_cell(row.cells[2], str(wscore), align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5)
    write_cell(row.cells[3], str(desc),   font_size_pt=10.5)

# 维度三：其他（scores[7]）
sr3 = tbl.add_row()
for i, c in enumerate(sr3.cells): set_cell_width(c, TM[i])
sc3 = sr3.cells[0].merge(sr3.cells[NC - 1])
write_cell(sc3, "三、其他", font_size_pt=10.5, bold=False)

add_col_header(tbl)

name, score, wscore, desc = SCORES[7]
row = tbl.add_row()
for i, c in enumerate(row.cells): set_cell_width(c, TM[i])
write_cell(row.cells[0], name,        font_size_pt=10.5)
write_cell(row.cells[1], str(score),  align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5)
write_cell(row.cells[2], str(wscore), align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5)
write_cell(row.cells[3], str(desc),   font_size_pt=10.5)

# 评价总分行
tr = tbl.add_row()
for i, c in enumerate(tr.cells): set_cell_width(c, TM[i])
lf = tr.cells[0].merge(tr.cells[1])
write_cell(lf, "评价总分",
           align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5, bold=True)
rf = tr.cells[2].merge(tr.cells[NC - 1])
if CONCL:
    concl_text = f"{TOTAL} 分    结论：{CONCL}"
else:
    concl_text = f"{TOTAL} 分"
write_cell(rf, concl_text,
           align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5, bold=True)

# ═══════════════════════════════════════════════════════════
# 第2页：附表1 主要问题与修改建议
# ═══════════════════════════════════════════════════════════
add_page_break(doc)
add_para(doc, "附表1  主要问题与修改建议",
         align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5, bold=True)

# 6列：序号 | 严重程度 | 问题名称 | 涉及位置 | 问题描述 | 修改建议
T2 = [1.0, 2.0, 3.0, 2.5, 4.25, 4.25]
t2 = doc.add_table(rows=0, cols=6)
set_col_widths(t2, T2)
set_borders(t2)

hr2 = t2.add_row()
for i, (c, w) in enumerate(zip(hr2.cells, T2)): set_cell_width(c, w)
for i, txt in enumerate(["序号", "严重程度", "问题名称", "涉及位置", "问题描述", "修改建议"]):
    write_cell(hr2.cells[i], txt,
               align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5, bold=True)

# issues 结构：(序号, 严重程度, 问题名称, 涉及位置, 问题描述, 修改建议)
for item in ISSUES:
    seq, level, name, loc, desc, sug = item
    dr = t2.add_row()
    for j, (c, w) in enumerate(zip(dr.cells, T2)): set_cell_width(c, w)
    write_cell(dr.cells[0], str(seq),  align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5)
    write_cell(dr.cells[1], level,     align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5)
    write_cell(dr.cells[2], name,      font_size_pt=10.5)
    write_cell(dr.cells[3], loc,       font_size_pt=10.5)
    write_cell(dr.cells[4], str(desc), font_size_pt=10.5)
    write_cell(dr.cells[5], str(sug),  font_size_pt=10.5)

# 补充空行（至少5行）
existing_data_rows = len(ISSUES)
for i in range(max(0, 5 - existing_data_rows)):
    dr = t2.add_row()
    for j, (c, w) in enumerate(zip(dr.cells, T2)): set_cell_width(c, w)
    write_cell(dr.cells[0], str(existing_data_rows + i + 1),
               align=WD_ALIGN_PARAGRAPH.CENTER, font_size_pt=10.5)
    for j in range(1, 6):
        write_cell(dr.cells[j], "", font_size_pt=10.5)

# 评价人信息行（3列合并）
ir = t2.add_row()
for j, (c, w) in enumerate(zip(ir.cells, T2)): set_cell_width(c, w)
# 评价人姓名（前2列）
m1 = ir.cells[0].merge(ir.cells[1])
write_cell(m1, f"评价人姓名：{RNAME}", font_size_pt=10.5)
# 评价日期（中间2列）
m2 = ir.cells[2].merge(ir.cells[3])
write_cell(m2, f"评价日期：{RDATE}", font_size_pt=10.5)
# 签字/盖章（后2列）
m3 = ir.cells[4].merge(ir.cells[5])
write_cell(m3, "签字/盖章：", font_size_pt=10.5)

# ── 保存 ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
abs_output_path = os.path.abspath(OUTPUT_PATH)
output_dir_path = os.path.dirname(abs_output_path)
word_filename = os.path.basename(abs_output_path)
word_session_path = ""
output_dir_session_path = ""
try:
    relative_output_path = os.path.relpath(abs_output_path, SESSION_ROOT)
    if not relative_output_path.startswith(".."):
        word_session_path = f"@session/{relative_output_path}"
        output_dir_session_path = f"@session/{os.path.dirname(relative_output_path)}"
except Exception:
    word_session_path = ""

open_output_dir = {
    "attempted": False,
    "command": "",
    "target": output_dir_path,
    "status": "skipped",
    "exit_code": None,
    "stderr": "",
}
if args.no_open_output:
    open_output_dir["status"] = "skipped_by_arg"
elif os.name == "nt":
    open_output_dir["attempted"] = True
    open_output_dir["command"] = "os.startfile"
    try:
        os.startfile(output_dir_path)
        open_output_dir["exit_code"] = 0
        open_output_dir["status"] = "completed"
    except Exception as exc:
        open_output_dir["status"] = "failed"
        open_output_dir["stderr"] = str(exc)
elif sys.platform == "darwin":
    open_output_dir["attempted"] = True
    open_output_dir["command"] = "open"
    try:
        open_result = subprocess.run(
            ["open", output_dir_path],
            check=False,
            capture_output=True,
            text=True,
            timeout=10,
        )
        open_output_dir["exit_code"] = open_result.returncode
        open_output_dir["stderr"] = (open_result.stderr or "").strip()
        open_output_dir["status"] = "completed" if open_result.returncode == 0 else "failed"
    except Exception as exc:
        open_output_dir["status"] = "failed"
        open_output_dir["stderr"] = str(exc)
else:
    open_output_dir["attempted"] = True
    open_output_dir["command"] = "xdg-open"
    try:
        open_result = subprocess.run(
            ["xdg-open", output_dir_path],
            check=False,
            capture_output=True,
            text=True,
            timeout=10,
        )
        open_output_dir["exit_code"] = open_result.returncode
        open_output_dir["stderr"] = (open_result.stderr or "").strip()
        open_output_dir["status"] = "completed" if open_result.returncode == 0 else "failed"
    except Exception as exc:
        open_output_dir["status"] = "failed"
        open_output_dir["stderr"] = str(exc)

if open_output_dir["status"] == "completed":
    final_response_text = "Word文件已生成，已自动打开输出目录。\n\n"
else:
    final_response_text = "Word文件已生成，但自动打开输出目录失败。\n\n"
final_response_text += f"输出目录：{output_dir_path}\n"
final_response_text += f"文件名：{word_filename}\n"
if word_session_path:
    final_response_text += f"会话文件：{word_session_path}\n"

final_response_path = os.path.join(
    output_dir_path,
    "final_response.md",
)
with open(final_response_path, "w", encoding="utf-8") as f:
    f.write(final_response_text)
manifest_path = os.path.join(output_dir_path, "word_output_manifest.json")
with open(manifest_path, "w", encoding="utf-8") as f:
    json.dump(
        {
            "kind": "word_output",
            "word_path": abs_output_path,
            "word_filename": word_filename,
            "word_session_path": word_session_path,
            "output_dir_path": output_dir_path,
            "output_dir_session_path": output_dir_session_path,
            "open_output_dir": open_output_dir,
            "final_response_path": final_response_path,
            "final_response_text": final_response_text,
            "data_path": DATA_FILE_PATH,
            "status": "generated",
        },
        f,
        ensure_ascii=False,
        indent=2,
    )
print(f"OK:{abs_output_path}")
print(f"OUTPUT_DIR:{output_dir_path}")
print(f"OPEN_OUTPUT_DIR:{open_output_dir['status']}")
if open_output_dir.get("stderr"):
    print(f"OPEN_OUTPUT_DIR_ERROR:{open_output_dir['stderr']}")
print(f"MANIFEST:{manifest_path}")
print(f"FINAL_RESPONSE:{final_response_path}")
print("FINAL_RESPONSE_BEGIN")
print(final_response_text.rstrip())
print("FINAL_RESPONSE_END")
