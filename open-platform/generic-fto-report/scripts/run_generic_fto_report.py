from __future__ import annotations

import argparse
import html
import json
import pathlib
import re
import time
from collections import OrderedDict
from datetime import datetime

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


SCRIPT_DIR = pathlib.Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
DEFAULT_API_CONFIG_PATH = SKILL_DIR / "references" / "zhihuiya_config.json"
DEFAULT_BUSINESS_CONFIG_PATH = SKILL_DIR / "references" / "config.json"

API_CONFIG_PATH = DEFAULT_API_CONFIG_PATH


def load_json(path: pathlib.Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8-sig"))


def load_api_config() -> tuple[str, str]:
    cfg = load_json(API_CONFIG_PATH)
    return cfg.get("zhihuiya_base_url", "https://connect.zhihuiya.com"), cfg["zhihuiya_api_key"]


def api_get(path: str, params: dict, *, timeout: int = 45) -> dict:
    base, apikey = load_api_config()
    resp = requests.get(base.rstrip("/") + path, params={"apikey": apikey, **params}, timeout=timeout)
    resp.raise_for_status()
    return resp.json()


def api_post(path: str, payload: dict, *, timeout: int = 60) -> dict | str:
    base, apikey = load_api_config()
    resp = requests.post(base.rstrip("/") + path, params={"apikey": apikey}, json=payload, timeout=timeout)
    resp.raise_for_status()
    try:
        return resp.json()
    except Exception:
        return resp.text


def split_values(text: str) -> list[str]:
    value = str(text or "").strip()
    for sep in ["，", "、", "；", ";", "\n", "|"]:
        value = value.replace(sep, ",")
    return [item.strip() for item in value.split(",") if item.strip()]


def first_nonempty(*values: str) -> str:
    for value in values:
        if str(value or "").strip():
            return str(value).strip()
    return ""


def normalize_metadata_key(key: str) -> str:
    key = re.sub(r"\s+", "", key or "").lower()
    aliases = {
        "产品名称": "product_name",
        "标的产品": "product_name",
        "项目名称": "project_name",
        "行业": "industry",
        "分类号": "ipc",
        "ipc": "ipc",
        "ipc分类号": "ipc",
        "主分类号": "mipc",
        "mipc": "mipc",
        "mipc分类号": "mipc",
        "竞争对手": "competitors",
        "竞争对手名称": "competitors",
        "排查竞争对手": "competitors",
        "排查对象": "competitors",
        "检索对象": "competitors",
        "申请人": "competitors",
        "专利权人": "competitors",
        "区域": "authority",
        "排查区域": "authority",
        "国家地区": "authority",
        "地区": "authority",
        "authority": "authority",
        "法律状态": "legal_status",
        "有效状态": "legal_status",
        "检索式": "queries",
        "自定义检索式": "queries",
        "风险点整体描述": "risk_summary",
        "技术方案": "risk_summary",
    }
    return aliases.get(key, key)


def normalize_ipc_expr(value: str) -> str:
    text = str(value or "").strip().strip("，,;；")
    for sep in ["、", "，", ",", "；", ";", "\n"]:
        text = text.replace(sep, " or ")
    text = re.sub(r"\s+or\s+", " or ", text, flags=re.I)
    return re.sub(r"\s+", " ", text).strip()


def normalize_authority(value: str) -> str:
    text = str(value or "").strip()
    mapping = {"中国": "CN", "中国大陆": "CN", "大陆": "CN", "美国": "US", "欧洲": "EP", "日本": "JP", "韩国": "KR"}
    return mapping.get(text, text)


def normalize_legal_status(value: str) -> str:
    text = str(value or "").strip()
    mapping = {"有效": "1", "有效专利": "1", "有效中": "1", "授权有效": "1", "失效": "0"}
    return mapping.get(text, text)


def parse_metadata_table(rows: list[list[str]]) -> dict:
    metadata: dict[str, str] = {}
    for row in rows:
        if len(row) < 2:
            continue
        key = normalize_metadata_key(row[0])
        value = row[1].strip()
        if key and value:
            if key == "queries" and key in metadata:
                metadata[key] = metadata[key] + "\n" + value
            else:
                metadata[key] = value
    return metadata


def parse_docx(path: pathlib.Path) -> tuple[list[dict], str, dict]:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    product_desc = "\n".join(paragraphs)
    metadata: dict[str, str] = {}
    features: list[dict] = []

    # Parse simple "字段：值" paragraphs as optional report/search metadata.
    for text in paragraphs:
        m = re.match(r"^([^:：]{2,24})[:：]\s*(.+)$", text)
        if m:
            key = normalize_metadata_key(m.group(1))
            value = m.group(2).strip()
            if key and value:
                metadata.setdefault(key, value)

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        header = rows[0]
        joined = "|".join(header)

        if ("技术特征" not in joined and "关键词" not in joined) or len(header) < 3:
            metadata.update({k: v for k, v in parse_metadata_table(rows).items() if v})
            continue

        type_idx, desc_idx, kw_idx = 0, 1, 2
        for i, h in enumerate(header):
            if "类型" in h or "序号" in h:
                type_idx = i
            elif "描述" in h or "内容" in h:
                desc_idx = i
            elif "关键词" in h or "关键字" in h:
                kw_idx = i
        for row in rows[1:]:
            if len(row) <= max(type_idx, desc_idx, kw_idx):
                continue
            type_text, desc_text, kw_text = row[type_idx], row[desc_idx], row[kw_idx]
            if not desc_text:
                continue
            important = any(token in type_text for token in ["重要", "核心", "必要", "关键"])
            keywords = split_values(kw_text)
            if not keywords:
                keywords = split_values(desc_text)[:3]
            features.append(
                {
                    "type": "important" if important else "normal",
                    "type_label": type_text or ("重要技术特征" if important else "技术特征"),
                    "description": desc_text,
                    "keywords": keywords,
                }
            )

    if not features:
        raise RuntimeError("未能从 Word 文档中解析到技术特征表")
    return features, product_desc, metadata


def expand_keywords(keyword: str, max_words: int = 12) -> list[str]:
    payload = {"keyword": [keyword], "lang": ["cn"], "type": ["synonym", "related"]}
    data = api_post("/search/patent/keyword-suggest", payload)
    out = [keyword]
    if isinstance(data, dict) and data.get("status") is True:
        for item in (data.get("data") or {}).get("items", []) or []:
            for obj in item.get("keyword_list", []) or []:
                kw = (obj.get("keyword") or "").strip()
                if kw and kw not in out:
                    out.append(kw)
    return out[:max_words]


def enrich_features(features: list[dict], *, skip_expansion: bool = False) -> list[dict]:
    for feat in features:
        expanded: list[str] = []
        for kw in feat["keywords"]:
            candidates = [kw] if skip_expansion else expand_keywords(kw)
            for item in candidates:
                if item and item not in expanded:
                    expanded.append(item)
        feat["expanded_keywords"] = expanded or feat["keywords"]
    return features


def quote_kw(word: str) -> str:
    return '"' + word.replace('"', '\\"') + '"'


def kw_clause(words: list[str]) -> str:
    return " OR ".join(quote_kw(w) for w in words if w)


def assignee_clause(competitors: list[str]) -> str:
    if not competitors:
        return ""
    return " AND ALL_AN:(" + " OR ".join(f'TREE@"{c}"' for c in competitors) + ")"


def build_suffix(scope: dict) -> str:
    parts: list[str] = []
    if scope.get("ipc"):
        parts.append(f"IPC:({scope['ipc']})")
    if scope.get("competitors"):
        parts.append(assignee_clause(scope["competitors"]).removeprefix(" AND "))
    if scope.get("legal_status"):
        parts.append(f"SIMPLE_LEGAL_STATUS:({scope['legal_status']})")
    if scope.get("authority"):
        parts.append(f"AUTHORITY:({scope['authority']})")
    return (" AND " + " AND ".join(parts)) if parts else ""


def load_manual_queries(path: pathlib.Path | None, inline_queries: list[str]) -> OrderedDict[str, str]:
    queries: OrderedDict[str, str] = OrderedDict()
    if path:
        raw = load_json(path)
        if isinstance(raw, dict):
            source = raw.get("queries", raw)
            if isinstance(source, dict):
                for name, query in source.items():
                    if str(query).strip():
                        queries[str(name)] = str(query).strip()
            elif isinstance(source, list):
                for i, query in enumerate(source, 1):
                    if str(query).strip():
                        queries[f"manual_query_{i}"] = str(query).strip()
    for i, query in enumerate(inline_queries, 1):
        if query.strip():
            queries[f"manual_query_{len(queries) + i}"] = query.strip()
    return queries


def parse_doc_manual_queries(metadata: dict) -> OrderedDict[str, str]:
    queries = OrderedDict()
    raw = str(metadata.get("queries") or "").strip()
    if not raw:
        return queries
    chunks = [c.strip() for c in re.split(r"\n\s*\n|(?:^|\n)\s*检索式\d*[:：]", raw) if c.strip()]
    if not chunks:
        chunks = [raw]
    for i, query in enumerate(chunks, 1):
        queries[f"doc_query_{i}"] = query
    return queries


def build_queries(features: list[dict], scope: dict, *, max_query2_subs: int = 12) -> OrderedDict[str, str]:
    important = [f for f in features if f["type"] == "important"] or features[:1]
    normal = [f for f in features if f["type"] != "important"]
    suffix = build_suffix(scope)

    q1_core = " AND ".join(f"({kw_clause(f['expanded_keywords'])})" for f in important)
    query1 = f"TAC_ALL:({q1_core}){suffix}"

    q2_parts: list[str] = []
    for imp in important:
        for norm in normal:
            q2_parts.append(f"TAC_ALL:(({kw_clause(imp['expanded_keywords'])}) AND ({kw_clause(norm['expanded_keywords'])})){suffix}")
            if len(q2_parts) >= max_query2_subs:
                break
        if len(q2_parts) >= max_query2_subs:
            break
    query2 = " OR ".join(f"({p})" for p in q2_parts) if q2_parts else query1

    query3 = ""
    if scope.get("mipc"):
        query3 = f"MIPC:({scope['mipc']})"
        if scope.get("competitors"):
            query3 += assignee_clause(scope["competitors"])
        if scope.get("legal_status"):
            query3 += f" AND SIMPLE_LEGAL_STATUS:({scope['legal_status']})"
        if scope.get("authority"):
            query3 += f" AND AUTHORITY:({scope['authority']})"

    queries = OrderedDict([("feature_all", query1), ("feature_pair", query2)])
    if query3:
        queries["classification"] = query3
    return queries


def search_patents(query: str, scope: dict, max_total: int) -> list[dict]:
    all_rows: list[dict] = []
    page_size = int(scope.get("page_size") or 100)
    offset = 0
    collapse_authority = split_values(scope.get("authority") or "") or ["CN", "US", "EP", "JP", "KR"]
    while offset < max_total:
        payload = {
            "query_text": query,
            "limit": min(page_size, max_total - offset),
            "offset": offset,
            "stemming": 0,
            "sort": [{"field": "SCORE", "order": "DESC"}],
            "collapse_by": "PBD",
            "collapse_type": "ALL",
            "collapse_order": "LATEST",
            "collapse_order_authority": collapse_authority,
        }
        data = api_post("/search/patent/query-search-patent/v2", payload, timeout=90)
        if not isinstance(data, dict) or data.get("status") is False:
            raise RuntimeError(f"P002 failed: {data}")
        results = (data.get("data") or {}).get("results", []) or []
        all_rows.extend(results)
        if len(results) < payload["limit"]:
            break
        offset += payload["limit"]
    return all_rows


def collect_patents(queries: OrderedDict[str, str], scope: dict) -> tuple[list[dict], dict[str, int]]:
    seen: set[str] = set()
    merged: list[dict] = []
    counts: dict[str, int] = {}
    for name, query in queries.items():
        rows = search_patents(query, scope, int(scope.get("max_total") or 500))
        counts[name] = len(rows)
        for row in rows:
            key = row.get("pn") or row.get("patent_id") or row.get("apno") or json.dumps(row, sort_keys=True, ensure_ascii=False)
            if key in seen:
                continue
            seen.add(key)
            row["matched_query"] = name
            merged.append(row)
    return merged, counts


def clean_claim_html(raw: str, claim_num: str = "1") -> str:
    parts = re.findall(rf'<div[^>]*num="{re.escape(claim_num)}"[^>]*>(.*?)</div>', raw, flags=re.S)
    if not parts:
        parts = [raw]
    text = "".join(re.sub(r"<[^>]+>", "", part) for part in parts)
    text = html.unescape(text)
    return re.sub(r"\s+", "", text).strip()


def get_claim1(pn: str) -> tuple[str, dict]:
    data = api_get("/basic-patent-data/claim-data", {"patent_number": pn, "replace_by_related": 0}, timeout=45)
    if data.get("status") is not True:
        return "", data
    rows = data.get("data") or []
    if not rows or not rows[0].get("claims"):
        return "", data
    raw = rows[0]["claims"][0].get("claim_text") or ""
    return clean_claim_html(raw, "1"), data


def feature_hit(feature: dict, claim_text: str) -> bool:
    words = feature.get("expanded_keywords") or feature.get("keywords") or []
    return any(w and w in claim_text for w in words)


def local_claim_compare(features: list[dict], claim1: str) -> dict:
    matched = []
    missing = []
    for feat in features:
        if feature_hit(feat, claim1):
            matched.append((feat["description"], "权利要求1中可见相同或近似关键词，需进一步核对结构/步骤/功能限定是否完全对应。"))
        else:
            missing.append(feat["description"])
    important = [f for f in features if f["type"] == "important"] or features
    important_hits = sum(1 for f in important if feature_hit(f, claim1))
    all_hits = sum(1 for f in features if feature_hit(f, claim1))
    important_ratio = important_hits / max(len(important), 1)
    all_ratio = all_hits / max(len(features), 1)
    if important_ratio >= 0.9 and all_ratio >= 0.65:
        risk = "High"
    elif important_ratio >= 0.6 or all_ratio >= 0.45:
        risk = "Medium"
    else:
        risk = "Low"
    return {
        "risk_level": risk,
        "matched_features": matched,
        "missing_or_diff_features": missing or ["未从关键词层面发现明显缺失，仍需人工核对权利要求限定。"],
        "conclusion": "该结论为基于权利要求1文本与输入技术特征的结构化初筛结论，正式 FTO 意见需结合完整权利要求、说明书、审查档案及产品资料复核。",
    }


def call_ai07(pn: str, title: str, claim1: str, features: list[dict]) -> dict:
    product = "；".join(f["description"] for f in features)[:220]
    prompt = (
        f"专利FTO侵权比对。专利{pn}《{title}》。"
        f"权利要求1摘要:{claim1[:220]}。"
        f"标的技术特征:{product}。"
        "请输出:字面侵权风险=High/Medium/Low;等同风险=High/Medium/Low;差异特征;一句报告结论。"
    )[:650]
    try:
        raw = api_post("/chat/cc-gpt-stream", {"prompt": prompt, "stream": True}, timeout=90)
        text = raw if isinstance(raw, str) else json.dumps(raw, ensure_ascii=False)
        chunks = []
        for m in re.finditer(r"data:(\{.*?\})(?=\s*data:|$)", text, flags=re.S):
            try:
                obj = json.loads(m.group(1))
                content = (((obj.get("data") or {}).get("output") or {}).get("content") or "")
                if content:
                    chunks.append(content)
            except Exception:
                pass
        content = "".join(chunks).strip()
        return {"ok": bool(content), "prompt": prompt, "content": content, "raw_excerpt": text[:1200]}
    except Exception as exc:
        return {"ok": False, "prompt": prompt, "content": "", "error": str(exc)}


def score_patent(pat: dict, features: list[dict]) -> int:
    text = " ".join(str(pat.get(k) or "") for k in ["title", "patsnap_title", "abstract", "current_assignee", "pn"])
    score = 0
    for feature in features:
        weight = 12 if feature["type"] == "important" else 6
        if any(kw and kw in text for kw in feature.get("expanded_keywords", [])):
            score += weight
    if str(pat.get("pn", "")).endswith("B"):
        score += 8
    if pat.get("matched_query") in {"feature_all", "manual_query_1", "doc_query_1"}:
        score += 12
    return score


def select_candidates(patents: list[dict], features: list[dict], limit: int) -> list[dict]:
    pool = sorted(((score_patent(p, features), p) for p in patents), key=lambda x: x[0], reverse=True)
    selected: list[dict] = []
    seen: set[str] = set()
    for _, pat in pool:
        pn = pat.get("pn") or pat.get("patent_id") or ""
        family_key = re.sub(r"[ABCU]\d?$", "", str(pn))
        if family_key in seen:
            continue
        seen.add(family_key)
        selected.append(pat)
        if len(selected) >= limit:
            break
    return selected


def fmt_date(value) -> str:
    s = str(value or "")
    if len(s) == 8 and s.isdigit():
        return f"{s[:4]}-{s[4:6]}-{s[6:]}"
    return s


def esc(value) -> str:
    return html.escape(str(value or ""))


def build_html_report(data: dict, output: pathlib.Path) -> None:
    scope = data["scope"]
    features = data["features"]
    candidates = data["candidates"]
    counts = data["query_counts"]
    risk_order = {"High": 0, "Medium": 1, "Low": 2, "Unknown": 3}
    sorted_candidates = sorted(candidates, key=lambda x: risk_order.get(x["comparison"]["risk_level"], 9))
    high_medium = [c for c in sorted_candidates if c["comparison"]["risk_level"] in {"High", "Medium"}]
    title = f"{scope.get('product_name') or scope.get('project_name') or '通用行业'} FTO 专利分析报告"
    feature_rows = "".join(
        f"<tr><td>{esc(f['type_label'])}</td><td>{esc(f['description'])}</td><td>{esc('、'.join(f['expanded_keywords']))}</td></tr>"
        for f in features
    )
    patent_rows = "".join(
        f"<tr><td>{i+1}</td><td>{esc(c.get('pn'))}</td><td>{esc(c.get('title') or c.get('patsnap_title'))}</td><td>{esc(c.get('current_assignee'))}</td>"
        f"<td>{fmt_date(c.get('apdt'))}</td><td>{fmt_date(c.get('pbdt'))}</td><td>{esc(c['comparison']['risk_level'])}</td></tr>"
        for i, c in enumerate(sorted_candidates)
    )
    query_blocks = "".join(
        f"<h3>{esc(k)}（P002返回 {counts.get(k, 0)} 条）</h3><pre>{esc(v)}</pre>" for k, v in data["queries"].items()
    )
    detail_blocks = ""
    for c in sorted_candidates:
        comp = c["comparison"]
        matched = "".join(f"<li>{esc(a)}：{esc(b)}</li>" for a, b in comp["matched_features"])
        diffs = "".join(f"<li>{esc(d)}</li>" for d in comp["missing_or_diff_features"])
        ai = c.get("ai07", {})
        ai_text = esc(ai.get("content") or ai.get("error") or "AI07未返回可结构化内容，原始响应已保存于JSON追溯文件。")
        detail_blocks += f"""
        <section class="card">
          <h3>{esc(c.get('pn'))} {esc(c.get('title') or c.get('patsnap_title'))} <span class="risk {esc(comp['risk_level'])}">{esc(comp['risk_level'])}</span></h3>
          <p><b>申请日：</b>{fmt_date(c.get('apdt'))}　<b>公开/公告日：</b>{fmt_date(c.get('pbdt'))}　<b>权利要求来源：</b>{esc(c.get('claim1_source'))}</p>
          <p><b>权利要求1摘录：</b>{esc(c.get('claim1', '')[:520])}{'...' if len(c.get('claim1', '')) > 520 else ''}</p>
          <p><b>命中特征：</b></p><ul>{matched}</ul>
          <p><b>差异/未覆盖特征：</b></p><ul>{diffs}</ul>
          <p><b>结构化结论：</b>{esc(comp['conclusion'])}</p>
          <p><b>AI07辅助记录：</b>{ai_text}</p>
        </section>
        """
    conclusion_text = (
        f"P002检索式合并去重后获得 {len(data['patent_list'])} 件专利记录，选取 {len(candidates)} 件重点关注专利获取P018权利要求1并进行比对。"
    )
    conclusion_text += f"其中 {len(high_medium)} 件被列为中/高关注。" if high_medium else "暂未发现关键词和权利要求1层面明显高相关的中/高风险专利。"
    html_doc = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><title>{esc(title)}</title>
<style>
body{{font-family:Arial,'Microsoft YaHei',sans-serif;margin:0;background:#f6f7f9;color:#1f2937;line-height:1.65}}
header{{background:#263238;color:white;padding:36px 52px}} header h1{{margin:0;font-size:28px}} header p{{margin:10px 0 0;color:#dfe7ea}}
main{{max-width:1120px;margin:24px auto;padding:0 20px}} section{{background:white;border:1px solid #e5e7eb;border-radius:6px;padding:22px;margin:16px 0}}
h2{{font-size:20px;border-bottom:2px solid #d7dee8;padding-bottom:8px;color:#263238}} h3{{font-size:16px;color:#263238}}
table{{width:100%;border-collapse:collapse;font-size:13px}} th,td{{border:1px solid #d8dee8;padding:8px;vertical-align:top}} th{{background:#eef2f7}}
pre{{white-space:pre-wrap;background:#f3f4f6;border:1px solid #e5e7eb;padding:10px;border-radius:4px;font-size:12px}}
.risk{{float:right;padding:2px 8px;border-radius:4px;color:white;font-size:12px}} .High{{background:#b91c1c}} .Medium{{background:#b45309}} .Low{{background:#15803d}} .Unknown{{background:#64748b}}
.note{{background:#fff7ed;border-color:#fed7aa}}
</style></head><body>
<header><h1>{esc(title)}</h1><p>生成日期：{datetime.now():%Y-%m-%d}</p></header>
<main>
<section><h2>结论摘要</h2><p>{esc(conclusion_text)}</p></section>
<section><h2>高相关/重点关注专利清单</h2><table><thead><tr><th>序号</th><th>公开(公告)号</th><th>名称</th><th>当前申请(专利权)人</th><th>申请日</th><th>公开日</th><th>风险等级</th></tr></thead><tbody>{patent_rows}</tbody></table></section>
<section><h2>一、分析目的</h2><p>本报告用于评估标的技术方案在指定国家/地区和竞争对手专利范围内的实施风险，为研发、上市、销售或出口决策提供专利风险分析参考。</p></section>
<section><h2>二、分析范围</h2><p>检索对象：{esc('、'.join(scope.get('competitors') or []) or '未限定')}；法律状态：{esc(scope.get('legal_status') or '未限定')}；地区：{esc(scope.get('authority') or '未限定')}；分类范围：IPC {esc(scope.get('ipc') or '未限定')}，MIPC {esc(scope.get('mipc') or '未限定')}。</p></section>
<section><h2>三、标的技术方案介绍</h2><p>{esc(data['product_summary'])}</p><table><thead><tr><th>类型</th><th>技术特征描述</th><th>关键词/扩展关键词</th></tr></thead><tbody>{feature_rows}</tbody></table></section>
<section><h2>四、检索策略</h2>{query_blocks}</section>
<section><h2>五、风险结论分析</h2>{detail_blocks}</section>
<section class="note"><h2>六、说明与免责</h2><p>本报告基于输入技术描述、智慧芽P070/P002/P018接口返回数据及AI07辅助调用结果生成。AI07仅作为辅助记录；正式法律意见仍需由专利律师结合完整权利要求、说明书、审查档案、法律状态和产品资料确认。</p></section>
</main></body></html>"""
    output.write_text(html_doc, encoding="utf-8")


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for i, text in enumerate(rows[0]):
        table.rows[0].cells[i].text = text
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
    return table


def build_docx_report(data: dict, output: pathlib.Path) -> None:
    scope = data["scope"]
    title = f"{scope.get('product_name') or scope.get('project_name') or '通用行业'} FTO 专利分析报告"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(38, 50, 56)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(38, 50, 56)
    doc.add_paragraph(f"生成日期：{datetime.now():%Y年%m月%d日}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_heading("结论摘要", level=1)
    doc.add_paragraph(f"P002检索式合并去重后获得 {len(data['patent_list'])} 件专利记录，选取 {len(data['candidates'])} 件重点关注专利获取P018权利要求1并进行比对。")
    doc.add_heading("表1：重点关注专利清单", level=1)
    rows = [["序号", "公开(公告)号", "名称", "申请日", "公开日", "风险等级"]]
    for i, c in enumerate(data["candidates"], 1):
        rows.append([str(i), c.get("pn") or "", c.get("title") or c.get("patsnap_title") or "", fmt_date(c.get("apdt")), fmt_date(c.get("pbdt")), c["comparison"]["risk_level"]])
    add_table(doc, rows)
    doc.add_heading("一、分析目的", level=1)
    doc.add_paragraph("本报告用于评估标的技术方案在指定国家/地区和竞争对手专利范围内的实施风险，为研发、上市、销售或出口决策提供专利风险分析参考。")
    doc.add_heading("二、分析范围", level=1)
    doc.add_paragraph(f"检索对象：{'、'.join(scope.get('competitors') or []) or '未限定'}；法律状态：{scope.get('legal_status') or '未限定'}；地区：{scope.get('authority') or '未限定'}；分类范围：IPC {scope.get('ipc') or '未限定'}，MIPC {scope.get('mipc') or '未限定'}。")
    doc.add_heading("三、标的技术方案介绍", level=1)
    doc.add_paragraph(data["product_summary"])
    rows = [["类型", "技术特征描述", "关键词/扩展关键词"]]
    for f in data["features"]:
        rows.append([f["type_label"], f["description"], "、".join(f["expanded_keywords"])])
    add_table(doc, rows)
    doc.add_heading("四、检索策略", level=1)
    for name, query in data["queries"].items():
        doc.add_heading(f"{name}（P002返回 {data['query_counts'].get(name, 0)} 条）", level=2)
        doc.add_paragraph(query)
    doc.add_heading("五、风险结论分析", level=1)
    for c in data["candidates"]:
        comp = c["comparison"]
        doc.add_heading(f"{c.get('pn')} {c.get('title') or c.get('patsnap_title') or ''}（{comp['risk_level']}）", level=2)
        doc.add_paragraph(f"申请日：{fmt_date(c.get('apdt'))}；公开/公告日：{fmt_date(c.get('pbdt'))}；权利要求来源：{c.get('claim1_source')}。")
        doc.add_paragraph("权利要求1摘录：" + c.get("claim1", "")[:700] + ("..." if len(c.get("claim1", "")) > 700 else ""))
        doc.add_paragraph("差异/未覆盖特征：")
        for d in comp["missing_or_diff_features"]:
            doc.add_paragraph(str(d))
        doc.add_paragraph("结论：" + comp["conclusion"])
    doc.add_heading("六、说明与免责", level=1)
    doc.add_paragraph("本报告基于输入技术描述、智慧芽P070/P002/P018接口返回数据及AI07辅助调用结果生成。AI07仅作为辅助记录；正式法律意见仍需由专利律师结合完整权利要求、说明书、审查档案、法律状态和产品资料确认。")
    doc.save(str(output))


def merge_scope(args, metadata: dict, config: dict) -> dict:
    scope = {
        "product_name": first_nonempty(args.product_name, metadata.get("product_name", ""), config.get("product_name", "")),
        "project_name": first_nonempty(args.project_name, metadata.get("project_name", ""), config.get("project_name", "")),
        "industry": first_nonempty(args.industry, metadata.get("industry", ""), config.get("industry", "")),
        "ipc": normalize_ipc_expr(first_nonempty(args.ipc, metadata.get("ipc", ""), config.get("ipc_filter", ""))),
        "mipc": normalize_ipc_expr(first_nonempty(args.mipc, metadata.get("mipc", ""), config.get("mipc_filter", ""))),
        "authority": normalize_authority(first_nonempty(args.authority, metadata.get("authority", ""), config.get("authority", ""))),
        "legal_status": normalize_legal_status(first_nonempty(args.legal_status, metadata.get("legal_status", ""), config.get("legal_status", ""))),
        "page_size": args.page_size or config.get("page_size", 100),
        "max_total": args.max_total or config.get("max_total", 500),
        "max_candidates": args.max_candidates or config.get("max_candidates", 10),
        "max_query2_subs": args.max_query2_subs or config.get("max_query2_subs", 12),
    }
    competitors = split_values(args.competitors) or split_values(metadata.get("competitors", "")) or split_values(config.get("target_company", ""))
    scope["competitors"] = competitors
    return scope


def main() -> None:
    global API_CONFIG_PATH
    parser = argparse.ArgumentParser(description="通用行业 FTO 报告生成器（skill 内部自洽版）")
    parser.add_argument("--docx", required=True, help="风险点描述 Word 文档路径")
    parser.add_argument("--output-dir", default="", help="输出目录，默认在当前目录下 generic_fto_report")
    parser.add_argument("--api-config", default=str(DEFAULT_API_CONFIG_PATH), help="skill 内部智慧芽 API 配置 JSON")
    parser.add_argument("--business-config", default=str(DEFAULT_BUSINESS_CONFIG_PATH), help="通用检索业务配置 JSON")
    parser.add_argument("--queries-json", default="", help="用户直接提供的检索式 JSON，格式可为 {\"queries\": {\"name\": \"query\"}}")
    parser.add_argument("--query", action="append", default=[], help="用户直接输入的一条检索式；可重复传入，至少需要提供 --query 或 --queries-json")
    parser.add_argument("--competitors", default="", help="竞争对手/申请人/专利权人，多个用逗号分隔")
    parser.add_argument("--ipc", default="", help="IPC 分类限制，例如 B60R21/ or B62D25/")
    parser.add_argument("--mipc", default="", help="主分类限制，用于 classification 检索式")
    parser.add_argument("--authority", default="", help="区域/专利局，例如 CN 或 US OR EP")
    parser.add_argument("--legal-status", default="", help="简单法律状态，例如 1 表示有效")
    parser.add_argument("--product-name", default="", help="标的产品/技术名称")
    parser.add_argument("--project-name", default="", help="项目名称")
    parser.add_argument("--industry", default="", help="行业名称")
    parser.add_argument("--skip-keyword-expansion", action="store_true", help="跳过 P070，仅使用输入关键词")
    parser.add_argument("--dry-run-queries", action="store_true", help="只解析输入并生成 queries.json/scope.json，不调用 P002/P018/AI07")
    parser.add_argument("--max-total", type=int, default=0, help="每条检索式最多拉取条数")
    parser.add_argument("--page-size", type=int, default=0, help="P002 分页大小")
    parser.add_argument("--max-candidates", type=int, default=0, help="进入 P018/AI07 比对的候选专利数量")
    parser.add_argument("--max-query2-subs", type=int, default=0, help="兼容旧版本参数；用户直接输入检索式时不使用")
    args = parser.parse_args()

    input_docx = pathlib.Path(args.docx)
    out_dir = pathlib.Path(args.output_dir) if args.output_dir else pathlib.Path.cwd() / "generic_fto_report"
    API_CONFIG_PATH = pathlib.Path(args.api_config)
    business_config = load_json(pathlib.Path(args.business_config))

    if not input_docx.exists():
        raise FileNotFoundError(input_docx)
    if not API_CONFIG_PATH.exists():
        raise FileNotFoundError(f"缺少智慧芽 API 配置：{API_CONFIG_PATH}")

    out_dir.mkdir(parents=True, exist_ok=True)
    features, product_desc, metadata = parse_docx(input_docx)
    scope = merge_scope(args, metadata, business_config)

    queries = load_manual_queries(pathlib.Path(args.queries_json) if args.queries_json else None, args.query)
    if not queries:
        raise RuntimeError("未检测到用户直接输入的检索式。请通过 --query \"检索式\" 传入，或通过 --queries-json 指定检索式 JSON。")
    query_mode = "manual"
    features = enrich_features(features, skip_expansion=args.skip_keyword_expansion)

    if args.dry_run_queries:
        dry_run_data = {
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "input_docx": str(input_docx),
            "query_mode": query_mode,
            "scope": scope,
            "metadata": metadata,
            "features": features,
            "product_summary": product_desc,
            "queries": queries,
        }
        (out_dir / "queries.json").write_text(json.dumps(queries, ensure_ascii=False, indent=2), encoding="utf-8")
        (out_dir / "scope.json").write_text(json.dumps(dry_run_data, ensure_ascii=False, indent=2), encoding="utf-8")
        print(str(out_dir))
        return

    patents, query_counts = collect_patents(queries, scope)
    selected = select_candidates(patents, features, limit=int(scope.get("max_candidates") or 10))
    enriched = []
    trace = {"p018": {}, "ai07": {}, "metadata": metadata}
    for pat in selected:
        pn = pat.get("pn") or ""
        claim1, raw = get_claim1(pn) if pn else ("", {"error": "missing pn"})
        trace["p018"][pn] = {"ok": bool(claim1), "raw_keys": list(raw.keys()) if isinstance(raw, dict) else []}
        pat["claim1"] = claim1
        pat["claim1_source"] = "p018_real" if claim1 else ("missing_identifier" if not pn else "p018_empty")
        pat["comparison"] = local_claim_compare(features, claim1) if claim1 else {
            "risk_level": "Unknown",
            "matched_features": [],
            "missing_or_diff_features": ["P018未返回权利要求1，无法形成侵权风险结论。"],
            "conclusion": "待补充权利要求后复核。",
        }
        pat["ai07"] = call_ai07(pn, pat.get("title") or pat.get("patsnap_title") or "", claim1, features) if claim1 else {"ok": False, "content": "", "error": "P018未返回权利要求1"}
        trace["ai07"][pn] = pat["ai07"]
        enriched.append(pat)
        time.sleep(0.2)

    data = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "input_docx": str(input_docx),
        "query_mode": query_mode,
        "scope": scope,
        "features": features,
        "product_summary": product_desc,
        "queries": queries,
        "query_counts": query_counts,
        "patent_list": patents,
        "candidates": enriched,
        "trace": trace,
    }
    (out_dir / "fto_structured_data.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    (out_dir / "patent_list.json").write_text(json.dumps(patents, ensure_ascii=False, indent=2), encoding="utf-8")
    (out_dir / "queries.json").write_text(json.dumps(queries, ensure_ascii=False, indent=2), encoding="utf-8")
    build_html_report(data, out_dir / "通用行业FTO专利分析报告.html")
    build_docx_report(data, out_dir / "通用行业FTO专利分析报告.docx")
    print(str(out_dir))


if __name__ == "__main__":
    main()
