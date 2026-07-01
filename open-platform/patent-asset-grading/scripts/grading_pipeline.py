#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
专利资产分级评审主流程脚本
用途：批量获取专利数据，按行业差异化权重评分，输出Excel/Word评审清单
依赖：openpyxl, python-docx
MCP依赖：patent-search（智慧芽）
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime, date
from typing import Optional

# ─────────────────────────────────────────────
# 1. 行业权重配置
# ─────────────────────────────────────────────

INDUSTRY_WEIGHTS = {
    "半导体/芯片":   {"D1": 0.35, "D2": 0.30, "D3": 0.15, "D4": 0.10, "D5": 0.10},
    "通信/ICT":     {"D1": 0.25, "D2": 0.20, "D3": 0.25, "D4": 0.10, "D5": 0.20},
    "医疗器械":     {"D1": 0.30, "D2": 0.30, "D3": 0.15, "D4": 0.15, "D5": 0.10},
    "生物医药":     {"D1": 0.25, "D2": 0.35, "D3": 0.20, "D4": 0.15, "D5": 0.05},
    "新能源/储能":  {"D1": 0.30, "D2": 0.25, "D3": 0.20, "D4": 0.15, "D5": 0.10},
    "消费电子":     {"D1": 0.25, "D2": 0.20, "D3": 0.30, "D4": 0.15, "D5": 0.10},
    "通用基准":     {"D1": 0.30, "D2": 0.25, "D3": 0.20, "D4": 0.15, "D5": 0.10},
}

# IPC前缀 → 行业映射（按顺序匹配，首次命中）
IPC_INDUSTRY_MAP = [
    (["H01L", "H01S", "G11C"],             "半导体/芯片"),
    (["H04W", "H04L", "H04B", "H04M"],     "通信/ICT"),
    (["A61B", "A61F", "A61M", "A61N"],     "医疗器械"),
    (["A61K", "A61P", "C07D", "C12N"],     "生物医药"),
    (["H01M", "H02J", "C01D", "C30B"],     "新能源/储能"),
    (["G06F", "G06T", "H04N", "G09G"],     "消费电子"),
]

GRADE_MAP = [
    (4.0, "S级", "战略核心资产"),
    (3.0, "A级", "重要资产"),
    (2.0, "B级", "一般资产"),
    (1.0, "C级", "补充资产"),
    (0.0, "D级", "待处置资产"),
]

DISPOSAL_MAP = {
    "S级": "重点维持，持续监控同族与侵权风险，优先纳入许可/诉讼策略",
    "A级": "正常维持，关注年费节点，评估许可或交叉许可机会",
    "B级": "按需维持，可作谈判筹码，超过8年剩余期考虑出售",
    "C级": "有限维持，建议评估放弃或打包出售，释放年费成本",
    "D级": "建议放弃或立即处置，避免无效年费支出",
}

# ─────────────────────────────────────────────
# 2. IPC → 行业识别
# ─────────────────────────────────────────────

def detect_industry(ipc_list: list[str], override: Optional[str] = None) -> str:
    """依据IPC列表识别所属行业；override可强制指定行业名称。"""
    if override:
        for ind in INDUSTRY_WEIGHTS:
            if override in ind:
                return ind
        return "通用基准"
    for ipc in ipc_list:
        code = ipc.strip().upper()
        for prefixes, industry in IPC_INDUSTRY_MAP:
            if any(code.startswith(p) for p in prefixes):
                return industry
    return "通用基准"


# ─────────────────────────────────────────────
# 3. 数据解析工具
# ─────────────────────────────────────────────

def parse_application_year(app_date_str: str) -> Optional[int]:
    """从申请日字符串解析年份（支持 YYYYMMDD / YYYY-MM-DD / YYYY）。"""
    if not app_date_str:
        return None
    app_date_str = str(app_date_str).strip()
    for fmt in ("%Y%m%d", "%Y-%m-%d", "%Y/%m/%d"):
        try:
            return datetime.strptime(app_date_str[:10], fmt[:len(app_date_str[:10])]).year
        except ValueError:
            pass
    try:
        return int(app_date_str[:4])
    except ValueError:
        return None


def calc_remaining_years(app_year: Optional[int], patent_type: str = "invention") -> float:
    """估算剩余保护期（年）。发明专利20年，实用新型10年，外观设计15年（CN最新）。"""
    if not app_year:
        return 0.0
    term_map = {"invention": 20, "utility": 10, "design": 15}
    term = term_map.get(patent_type, 20)
    expiry_year = app_year + term
    remaining = expiry_year - date.today().year
    return max(0.0, float(remaining))


def remaining_years_to_score(years: float) -> tuple[int, str]:
    """剩余保护期 → D4分值及依据。"""
    if years >= 15:
        return 5, f"剩余保护期约{years:.0f}年（≥15年），保护窗口充足"
    elif years >= 10:
        return 4, f"剩余保护期约{years:.0f}年（10–14年），中长期保护有效"
    elif years >= 5:
        return 3, f"剩余保护期约{years:.0f}年（5–9年），保护期尚可"
    elif years >= 2:
        return 2, f"剩余保护期约{years:.0f}年（2–4年），接近届满需关注"
    else:
        return 1, f"剩余保护期不足2年（约{years:.1f}年），即将届满或已届满"


def family_count_to_d3(family_count: int, key_markets: list[str]) -> tuple[int, str]:
    """同族数量+重要市场覆盖 → D3分值及依据。"""
    key = [m for m in key_markets if m.upper() in ("CN", "US", "EP", "JP", "WO", "KR")]
    key_str = "/".join(key) if key else "无主要市场"
    if family_count >= 8 or len(key) >= 4:
        return 5, f"同族{family_count}件，覆盖{key_str}等{len(key)}个主要市场，全球布局完善"
    elif family_count >= 5 or len(key) >= 3:
        return 4, f"同族{family_count}件，覆盖{key_str}，主要市场覆盖较好"
    elif family_count >= 3 or len(key) >= 2:
        return 3, f"同族{family_count}件，覆盖{key_str}，市场覆盖适中"
    elif family_count >= 2 or len(key) >= 1:
        return 2, f"同族{family_count}件，覆盖{key_str}，布局较为有限"
    else:
        return 1, f"仅单一国家申请（{key_str}），无多国布局"


def citation_to_d5(cited_count: int, legal_status: str, title: str, abstract: str) -> tuple[int, str]:
    """被引次数+法律状态 → D5分值及依据。"""
    notes = []
    score = 1

    if cited_count >= 50:
        score = 5
        notes.append(f"被引{cited_count}次，高影响力")
    elif cited_count >= 20:
        score = 4
        notes.append(f"被引{cited_count}次，引用较多")
    elif cited_count >= 10:
        score = 3
        notes.append(f"被引{cited_count}次，有一定引用")
    elif cited_count >= 3:
        score = 2
        notes.append(f"被引{cited_count}次，引用较少")
    else:
        score = 1
        notes.append(f"被引{cited_count}次，几乎无引用记录")

    # 法律状态附加信息
    status_lower = legal_status.lower() if legal_status else ""
    if "license" in status_lower or "许可" in legal_status:
        notes.append("存在许可记录")
        score = min(5, score + 1)
    if "litigation" in status_lower or "诉讼" in legal_status:
        notes.append("涉及诉讼")
        score = min(5, score + 1)

    # SEP关键词检测（标题/摘要）
    sep_keywords = ["standard", "3GPP", "IEEE", "SEP", "标准必要", "必要专利", "FRAND"]
    text = (title or "") + " " + (abstract or "")
    if any(kw.lower() in text.lower() for kw in sep_keywords):
        notes.append("疑似SEP/标准相关专利")
        score = min(5, score + 1)

    return score, "；".join(notes)

# ─────────────────────────────────────────────
# 4. AI评分函数（D1技术关联性 + D2权利要求强度）
# ─────────────────────────────────────────────

def score_d1_tech_relevance(title: str, abstract: str, claims: str, ipc_list: list[str], industry: str) -> tuple[int, str]:
    """
    D1技术关联性：基于IPC分类与行业核心技术方向的契合度判断。
    在Skill运行时，此函数由Eureka AI基于专利文本上下文推断打分。
    此处提供规则化回退逻辑（供无AI上下文时使用）。
    """
    # 规则回退：IPC与行业的匹配精准度作为代理
    industry_ipc_cores = {
        "半导体/芯片":  ["H01L", "H01S", "G11C", "H01J"],
        "通信/ICT":    ["H04W", "H04L", "H04B", "H04M", "H04Q"],
        "医疗器械":    ["A61B", "A61F", "A61M", "A61N", "A61H"],
        "生物医药":    ["A61K", "A61P", "C07D", "C07K", "C12N", "C12Q"],
        "新能源/储能": ["H01M", "H02J", "H02S", "C01D", "C30B"],
        "消费电子":    ["G06F", "G06T", "H04N", "G09G", "G06N"],
        "通用基准":    [],
    }
    core_ipc = industry_ipc_cores.get(industry, [])
    matched = [ipc for ipc in ipc_list if any(ipc.startswith(c) for c in core_ipc)]
    ratio = len(matched) / max(len(ipc_list), 1)

    # 权利要求数量作为辅助（特征越少越聚焦）
    claim_count = len(re.findall(r'(?:^|\n)\s*\d+[\.\、]', claims)) if claims else 0
    focus_bonus = 1 if 0 < claim_count <= 15 else 0

    if ratio >= 0.8:
        score = min(5, 4 + focus_bonus)
        reason = f"IPC分类与{industry}核心技术高度契合（{len(matched)}/{len(ipc_list)}个IPC命中），技术定向集中"
    elif ratio >= 0.5:
        score = 3
        reason = f"IPC分类与{industry}核心技术基本契合（{len(matched)}/{len(ipc_list)}个IPC命中），有一定关联"
    elif ratio >= 0.2:
        score = 2
        reason = f"IPC分类与{industry}核心技术部分重叠（{len(matched)}/{len(ipc_list)}个IPC命中），关联性一般"
    else:
        score = 1
        reason = f"IPC分类与{industry}核心技术领域关联性弱，技术方向偏离"
    return score, reason


def score_d2_claim_strength(claims: str) -> tuple[int, str]:
    """D2权利要求强度：基于独立权利要求数量与技术特征数估算保护范围。"""
    if not claims:
        return 2, "未获取权利要求文本，按默认值处理"

    # 识别独立权利要求（不含"根据权利要求N"）
    all_claims = re.split(r'\n\s*\d+[\.\、]', claims)
    independent = [c for c in all_claims if not re.search(r'根据权利要求\s*\d+|according to claim\s+\d+', c, re.IGNORECASE)]
    indep_count = max(1, len(independent))

    # 第一独立权利要求技术特征数（逗号分句估算）
    first_claim = independent[0] if independent else (all_claims[0] if all_claims else "")
    feature_count = len(re.split(r'[，,；;]', first_claim)) if first_claim else 10

    reasons = []
    if indep_count <= 3:
        reasons.append(f"{indep_count}个独立权项，独权数量精简")
    else:
        reasons.append(f"{indep_count}个独立权项，数量偏多，需关注单权项强度")

    if feature_count <= 5:
        score = 5
        reasons.append(f"第一独立权利要求约{feature_count}个技术特征，保护范围宽")
    elif feature_count <= 8:
        score = 4
        reasons.append(f"第一独立权利要求约{feature_count}个技术特征，保护范围较宽")
    elif feature_count <= 12:
        score = 3
        reasons.append(f"第一独立权利要求约{feature_count}个技术特征，保护范围适中")
    elif feature_count <= 18:
        score = 2
        reasons.append(f"第一独立权利要求约{feature_count}个技术特征，保护范围偏窄")
    else:
        score = 1
        reasons.append(f"第一独立权利要求约{feature_count}个技术特征，保护范围极窄，可规避风险高")

    if indep_count > 3:
        score = max(1, score - 1)

    return score, "；".join(reasons)


# ─────────────────────────────────────────────
# 5. 附加修正项
# ─────────────────────────────────────────────

def calc_bonus(legal_status: str, title: str, abstract: str) -> tuple[float, str]:
    """计算附加修正分值。"""
    bonus = 0.0
    reasons = []
    text = (title or "") + " " + (abstract or "") + " " + (legal_status or "")

    sep_kws = ["standard", "3GPP", "IEEE", "SEP", "标准必要", "必要专利", "ETSI", "ITU"]
    if any(kw.lower() in text.lower() for kw in sep_kws):
        bonus += 0.5
        reasons.append("疑似SEP（+0.5）")

    litigation_kws = ["litigation", "infringement", "诉讼", "侵权", "337"]
    if any(kw.lower() in text.lower() for kw in litigation_kws):
        if "有效" in text or "维持" in text or "valid" in text.lower():
            bonus += 0.3
            reasons.append("涉诉有效维持（+0.3）")

    invalid_kws = ["inter partes review", "IPR", "无效请求", "无效宣告", "复审", "PTR", "opposition"]
    if any(kw.lower() in text.lower() for kw in invalid_kws):
        bonus -= 0.5
        reasons.append("存在无效/复审风险（-0.5）")

    terminated_kws = ["abandoned", "lapsed", "expired", "withdrawn", "放弃", "失效", "终止", "已届满"]
    if any(kw.lower() in text.lower() for kw in terminated_kws):
        bonus -= 0.3
        reasons.append("权利已终止或产品停产（-0.3）")

    return bonus, "；".join(reasons) if reasons else "无附加修正"


# ─────────────────────────────────────────────
# 6. 等级判定
# ─────────────────────────────────────────────

def calc_grade(weighted_score: float) -> tuple[str, str]:
    for threshold, grade, name in GRADE_MAP:
        if weighted_score >= threshold:
            return grade, name
    return "D级", "待处置资产"

# ─────────────────────────────────────────────
# 7. 核心评审函数：评审单件专利（接收结构化数据）
# ─────────────────────────────────────────────

def grade_patent(patent_data: dict, industry_override: Optional[str] = None) -> dict:
    """
    对单件专利执行完整评审。
    patent_data 字段（均为可选，缺失时降级处理）：
      - pn: 专利号
      - title: 标题
      - abstract: 摘要
      - claims: 权利要求文本
      - ipc: IPC分类列表
      - application_date: 申请日（YYYYMMDD或YYYY-MM-DD）
      - legal_status: 法律状态文本
      - family_count: 同族数量（int）
      - family_jurisdictions: 同族国家代码列表
      - cited_count: 被引次数（int）
      - assignee: 申请人/专利权人
      - patent_type: invention/utility/design（默认invention）
    """
    pn = patent_data.get("pn", "N/A")
    title = patent_data.get("title", "")
    abstract = patent_data.get("abstract", "")
    claims = patent_data.get("claims", "")
    ipc_list = patent_data.get("ipc", [])
    app_date = patent_data.get("application_date", "")
    legal_status = patent_data.get("legal_status", "")
    family_count = int(patent_data.get("family_count", 1))
    family_jurisdictions = patent_data.get("family_jurisdictions", [])
    cited_count = int(patent_data.get("cited_count", 0))
    patent_type = patent_data.get("patent_type", "invention")
    assignee = patent_data.get("assignee", "")

    # 行业识别
    industry = detect_industry(ipc_list, industry_override)
    weights = INDUSTRY_WEIGHTS[industry]

    # 各维度评分
    d1_score, d1_reason = score_d1_tech_relevance(title, abstract, claims, ipc_list, industry)
    d2_score, d2_reason = score_d2_claim_strength(claims)

    app_year = parse_application_year(app_date)
    remaining = calc_remaining_years(app_year, patent_type)
    d4_score, d4_reason = remaining_years_to_score(remaining)
    d3_score, d3_reason = family_count_to_d3(family_count, family_jurisdictions)
    d5_score, d5_reason = citation_to_d5(cited_count, legal_status, title, abstract)

    # 加权总分
    raw_score = (
        d1_score * weights["D1"] +
        d2_score * weights["D2"] +
        d3_score * weights["D3"] +
        d4_score * weights["D4"] +
        d5_score * weights["D5"]
    )

    # 附加修正
    bonus, bonus_reason = calc_bonus(legal_status, title, abstract)
    final_score = round(min(5.0, max(0.0, raw_score + bonus)), 2)

    grade, grade_name = calc_grade(final_score)
    disposal = DISPOSAL_MAP.get(grade, "")

    return {
        "序号": None,  # 由批量函数填充
        "专利号": pn,
        "标题": title[:80] if title else "",
        "申请人": assignee,
        "申请日": app_date,
        "法律状态": legal_status,
        "所属行业": industry,
        "D1_技术关联性_分值": d1_score,
        "D1_技术关联性_依据": d1_reason,
        "D2_权利要求强度_分值": d2_score,
        "D2_权利要求强度_依据": d2_reason,
        "D3_市场覆盖度_分值": d3_score,
        "D3_市场覆盖度_依据": d3_reason,
        "D4_剩余保护期_分值": d4_score,
        "D4_剩余保护期_依据": d4_reason,
        "D5_被引用交叉价值_分值": d5_score,
        "D5_被引用交叉价值_依据": d5_reason,
        "附加修正": round(bonus, 1),
        "附加修正说明": bonus_reason,
        "加权总分": final_score,
        "最终等级": grade,
        "等级名称": grade_name,
        "处置建议": disposal,
    }


# ─────────────────────────────────────────────
# 8. 数据获取：解析MCP返回的Markdown文本
# ─────────────────────────────────────────────

def parse_patent_markdown(pn: str, markdown: str) -> dict:
    """
    从智慧芽MCP返回的Markdown文本中提取结构化字段。
    此函数在Skill运行时由Eureka AI调用，这里提供正则化回退提取逻辑。
    """
    data = {"pn": pn}

    def extract(pattern, text, group=1, default=""):
        m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        return m.group(group).strip() if m else default

    # 标题
    data["title"] = extract(r'(?:Title|标题)[：:]\s*(.+?)(?:\n|$)', markdown) or \
                    extract(r'^#\s+(.+?)$', markdown, flags=re.MULTILINE)

    # 摘要
    data["abstract"] = extract(r'(?:Abstract|摘要)[：:]\s*([\s\S]+?)(?:\n#+|\Z)', markdown)[:800]

    # 申请人
    data["assignee"] = extract(r'(?:Assignee|申请人|专利权人)[：:]\s*(.+?)(?:\n|$)', markdown)

    # 申请日
    data["application_date"] = extract(r'(?:Application Date|申请日)[：:]\s*(\d{4}[-/]?\d{2}[-/]?\d{2})', markdown)

    # IPC
    ipc_raw = extract(r'(?:IPC|国际分类号)[：:]\s*(.+?)(?:\n|$)', markdown)
    data["ipc"] = [i.strip() for i in re.split(r'[;；,，\s]+', ipc_raw) if re.match(r'[A-H]\d{2}', i.strip())] if ipc_raw else []

    # 法律状态
    data["legal_status"] = extract(r'(?:Legal Status|法律状态)[：:]\s*(.+?)(?:\n|$)', markdown)

    # 权利要求（取前3000字符）
    claims_raw = extract(r'(?:Claims|权利要求书?)[：:]\s*([\s\S]+?)(?:\n#+|\Z)', markdown)
    data["claims"] = claims_raw[:3000] if claims_raw else ""

    # 同族数量
    family_str = extract(r'(?:Family|同族)[^:：\n]*[：:]\s*(\d+)', markdown)
    data["family_count"] = int(family_str) if family_str.isdigit() else 1

    # 同族国家
    jurisdictions_raw = extract(r'(?:Jurisdictions|布局国家|同族国家)[：:]\s*(.+?)(?:\n|$)', markdown)
    data["family_jurisdictions"] = [j.strip().upper() for j in re.split(r'[,，\s/]+', jurisdictions_raw) if len(j.strip()) == 2] if jurisdictions_raw else []

    # 被引次数
    cited_str = extract(r'(?:Cited By|被引次数|被引频次)[：:]\s*(\d+)', markdown)
    data["cited_count"] = int(cited_str) if cited_str.isdigit() else 0

    # 专利类型推断（从专利号判断）
    if re.search(r'ZL?\d+\.\d|utility|实用新型', pn + markdown, re.IGNORECASE):
        data["patent_type"] = "utility"
    elif re.search(r'design|外观|ZL?\d+\.\d{4}D', pn + markdown, re.IGNORECASE):
        data["patent_type"] = "design"
    else:
        data["patent_type"] = "invention"

    return data

# ─────────────────────────────────────────────
# 9. Excel输出
# ─────────────────────────────────────────────

def export_excel(results: list[dict], output_path: str) -> str:
    """将评审结果输出为格式化Excel文件。"""
    try:
        import openpyxl
        from openpyxl.styles import (Font, PatternFill, Alignment,
                                      Border, Side, numbers)
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("[ERROR] 缺少 openpyxl，请执行：pip install openpyxl", file=sys.stderr)
        sys.exit(1)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "专利资产分级评审清单"

    # 颜色定义
    GRADE_COLORS = {
        "S级": "FF2E75B6",  # 深蓝
        "A级": "FF00B050",  # 绿
        "B级": "FFFFF2CC",  # 浅黄（字深）
        "C级": "FFFFD966",  # 黄
        "D级": "FFFF0000",  # 红
    }
    HEADER_FILL = PatternFill("solid", fgColor="FF1F3864")
    SUBHEADER_FILL = PatternFill("solid", fgColor="FF2E75B6")
    thin = Side(style="thin", color="FFAAAAAA")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── 标题行 ──
    ws.merge_cells("A1:V1")
    title_cell = ws["A1"]
    title_cell.value = "专利资产分级评审结果清单"
    title_cell.font = Font(name="微软雅黑", bold=True, size=14, color="FFFFFFFF")
    title_cell.fill = HEADER_FILL
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    # ── 生成时间行 ──
    ws.merge_cells("A2:V2")
    ws["A2"].value = f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}  |  共 {len(results)} 件专利"
    ws["A2"].font = Font(name="微软雅黑", size=9, color="FF666666")
    ws["A2"].alignment = Alignment(horizontal="right")
    ws.row_dimensions[2].height = 16

    # ── 表头 ──
    headers = [
        ("序号", 6),
        ("专利号", 16),
        ("标题", 36),
        ("申请人", 18),
        ("申请日", 12),
        ("法律状态", 10),
        ("所属行业", 12),
        ("D1\n技术关联性\n分值", 8),
        ("D1依据", 30),
        ("D2\n权利要求强度\n分值", 8),
        ("D2依据", 30),
        ("D3\n市场覆盖度\n分值", 8),
        ("D3依据", 30),
        ("D4\n剩余保护期\n分值", 8),
        ("D4依据", 24),
        ("D5\n被引/交叉\n分值", 8),
        ("D5依据", 24),
        ("附加修正", 8),
        ("附加修正说明", 20),
        ("加权总分", 9),
        ("最终等级", 8),
        ("处置建议", 30),
    ]

    for col_idx, (header, width) in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=col_idx, value=header)
        cell.font = Font(name="微软雅黑", bold=True, size=9, color="FFFFFFFF")
        cell.fill = SUBHEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    ws.row_dimensions[3].height = 40

    # ── 数据行 ──
    data_keys = [
        "序号", "专利号", "标题", "申请人", "申请日", "法律状态", "所属行业",
        "D1_技术关联性_分值", "D1_技术关联性_依据",
        "D2_权利要求强度_分值", "D2_权利要求强度_依据",
        "D3_市场覆盖度_分值", "D3_市场覆盖度_依据",
        "D4_剩余保护期_分值", "D4_剩余保护期_依据",
        "D5_被引用交叉价值_分值", "D5_被引用交叉价值_依据",
        "附加修正", "附加修正说明", "加权总分", "最终等级", "处置建议"
    ]

    for row_idx, result in enumerate(results, start=1):
        result["序号"] = row_idx
        r = row_idx + 3
        grade = result.get("最终等级", "")
        row_fill_color = GRADE_COLORS.get(grade, "FFFFFFFF")

        for col_idx, key in enumerate(data_keys, start=1):
            val = result.get(key, "")
            cell = ws.cell(row=r, column=col_idx, value=val)
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = Font(name="微软雅黑", size=9)

            # 分值列居中
            if "分值" in key or key in ("序号", "附加修正", "加权总分"):
                cell.alignment = Alignment(horizontal="center", vertical="top")

            # 等级列着色
            if key == "最终等级":
                cell.fill = PatternFill("solid", fgColor=row_fill_color)
                cell.font = Font(name="微软雅黑", size=9, bold=True,
                                 color="FFFFFFFF" if grade in ("S级", "A级", "D级") else "FF333333")
                cell.alignment = Alignment(horizontal="center", vertical="top")

            # 加权总分两位小数
            if key == "加权总分":
                cell.number_format = "0.00"

        ws.row_dimensions[r].height = 48

    # 冻结前3行和前2列
    ws.freeze_panes = "C4"

    # ── 说明sheet ──
    ws2 = wb.create_sheet("评审标准说明")
    notes = [
        ["专利资产分级评审标准说明", ""],
        ["", ""],
        ["五大评分维度（每项1–5分）", ""],
        ["D1 技术关联性", "评估IPC分类与行业核心技术方向的契合度（权重因行业而异）"],
        ["D2 权利要求强度", "评估独立权利要求数量与技术特征数，衡量保护范围宽窄"],
        ["D3 市场覆盖度", "同族专利覆盖国家/地区数量，重点关注CN/US/EP/JP/WO"],
        ["D4 剩余保护期", "≥15年=5分；10–14年=4分；5–9年=3分；2–4年=2分；<2年=1分"],
        ["D5 被引/交叉价值", "被引次数及是否涉及SEP/标准/许可/诉讼"],
        ["", ""],
        ["等级映射", ""],
        ["S级（战略核心）", "加权总分 ≥ 4.0"],
        ["A级（重要资产）", "加权总分 3.0–3.9"],
        ["B级（一般资产）", "加权总分 2.0–2.9"],
        ["C级（补充资产）", "加权总分 1.0–1.9"],
        ["D级（待处置）",   "加权总分 < 1.0"],
        ["", ""],
        ["附加修正项", ""],
        ["SEP声明", "+0.5"],
        ["涉诉有效维持", "+0.3"],
        ["高无效风险", "-0.5"],
        ["对应产品已停产/权利终止", "-0.3"],
    ]
    for row in notes:
        ws2.append(row)
    ws2.column_dimensions["A"].width = 24
    ws2.column_dimensions["B"].width = 60

    wb.save(output_path)
    return output_path


# ─────────────────────────────────────────────
# 10. Word输出
# ─────────────────────────────────────────────

def export_word(results: list[dict], output_path: str) -> str:
    """将评审结果输出为Word格式，每件专利一张评审卡片。"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[ERROR] 缺少 python-docx，请执行：pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面标题
    title_para = doc.add_heading("专利资产分级评审结果清单", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}  |  共 {len(results)} 件专利")
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    GRADE_COLORS_RGB = {
        "S级": RGBColor(0x1F, 0x37, 0x64),
        "A级": RGBColor(0x00, 0xB0, 0x50),
        "B级": RGBColor(0xFF, 0xD9, 0x66),
        "C级": RGBColor(0xFF, 0x92, 0x00),
        "D级": RGBColor(0xFF, 0x00, 0x00),
    }

    for i, result in enumerate(results, start=1):
        # 专利卡片标题
        h = doc.add_heading(f"[{i}] {result.get('最终等级','')} — {result.get('专利号','')}", level=2)

        # 基本信息表
        info_table = doc.add_table(rows=2, cols=4)
        info_table.style = "Table Grid"
        labels = ["专利号", "申请人", "申请日", "法律状态"]
        values = [result.get(k, "") for k in labels]
        for j, (lbl, val) in enumerate(zip(labels, values)):
            cell_l = info_table.cell(0, j)
            cell_v = info_table.cell(1, j)
            cell_l.text = lbl
            cell_l.paragraphs[0].runs[0].font.bold = True
            cell_v.text = str(val)

        doc.add_paragraph()

        # 五维评分明细表
        score_table = doc.add_table(rows=7, cols=3)
        score_table.style = "Table Grid"
        score_headers = ["评分维度", "分值（1–5）", "评分依据"]
        for j, hdr in enumerate(score_headers):
            c = score_table.cell(0, j)
            c.text = hdr
            c.paragraphs[0].runs[0].font.bold = True

        dim_rows = [
            ("D1 技术关联性", result.get("D1_技术关联性_分值",""), result.get("D1_技术关联性_依据","")),
            ("D2 权利要求强度", result.get("D2_权利要求强度_分值",""), result.get("D2_权利要求强度_依据","")),
            ("D3 市场覆盖度", result.get("D3_市场覆盖度_分值",""), result.get("D3_市场覆盖度_依据","")),
            ("D4 剩余保护期", result.get("D4_剩余保护期_分值",""), result.get("D4_剩余保护期_依据","")),
            ("D5 被引/交叉价值", result.get("D5_被引用交叉价值_分值",""), result.get("D5_被引用交叉价值_依据","")),
            ("附加修正", result.get("附加修正",""), result.get("附加修正说明","")),
        ]
        for row_j, (dim, sc, reason) in enumerate(dim_rows, start=1):
            score_table.cell(row_j, 0).text = dim
            score_table.cell(row_j, 1).text = str(sc)
            score_table.cell(row_j, 2).text = str(reason)

        doc.add_paragraph()

        # 结论段
        conclusion = doc.add_paragraph()
        conclusion.add_run("最终等级：").bold = True
        grade_run = conclusion.add_run(f"{result.get('最终等级','')}（{result.get('等级名称','')}）")
        grade_run.bold = True
        color = GRADE_COLORS_RGB.get(result.get("最终等级",""), RGBColor(0,0,0))
        grade_run.font.color.rgb = color

        conclusion.add_run(f"  加权总分：{result.get('加权总分',0):.2f}")

        disposal_para = doc.add_paragraph()
        disposal_para.add_run("处置建议：").bold = True
        disposal_para.add_run(result.get("处置建议", ""))

        industry_para = doc.add_paragraph()
        industry_para.add_run("所属行业：").bold = True
        industry_para.add_run(result.get("所属行业", ""))

        # 分隔线（非最后一件）
        if i < len(results):
            doc.add_paragraph("─" * 60)

    doc.save(output_path)
    return output_path

# ─────────────────────────────────────────────
# 11. 主流程入口（供Eureka Skill调用）
# ─────────────────────────────────────────────

def run_grading(
    patent_numbers: list[str],
    patent_markdown_map: dict[str, str],
    output_format: str = "excel",
    output_dir: str = ".",
    industry_override: Optional[str] = None,
) -> str:
    """
    主评审流程。
    
    参数：
      patent_numbers      - 专利号列表
      patent_markdown_map - {专利号: 智慧芽MCP返回的Markdown文本}，由Eureka在调用前填充
      output_format       - "excel" 或 "word"
      output_dir          - 输出目录
      industry_override   - 可选，强制指定行业（如"医疗器械"）
    
    返回：
      输出文件路径
    """
    results = []
    for pn in patent_numbers:
        md = patent_markdown_map.get(pn, "")
        if not md:
            print(f"[WARN] {pn}: 未获取到数据，以空数据评审", file=sys.stderr)
        patent_data = parse_patent_markdown(pn, md)
        result = grade_patent(patent_data, industry_override)
        results.append(result)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    os.makedirs(output_dir, exist_ok=True)

    if output_format.lower() == "word":
        output_path = os.path.join(output_dir, f"专利资产分级评审清单_{timestamp}.docx")
        return export_word(results, output_path)
    else:
        output_path = os.path.join(output_dir, f"专利资产分级评审清单_{timestamp}.xlsx")
        return export_excel(results, output_path)


# ─────────────────────────────────────────────
# 12. CLI入口（供测试及直接调用）
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="专利资产分级评审工具 — CLI入口",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例：
  python grading_pipeline.py --patents CN202310001234A CN114567890B --output-format excel
  python grading_pipeline.py --patents-file patents.txt --output-format word --industry 医疗器械
  python grading_pipeline.py --patents CN202310001234A --mock-data

patents.txt 格式：每行一个专利号（支持#注释行）
        """
    )
    parser.add_argument("--patents", nargs="+", metavar="PN",
                        help="专利号列表（空格分隔）")
    parser.add_argument("--patents-file", metavar="FILE",
                        help="包含专利号的文本文件（每行一个）")
    parser.add_argument("--output-format", choices=["excel", "word"], default="excel",
                        help="输出格式（默认excel）")
    parser.add_argument("--output-dir", default=".",
                        help="输出目录（默认当前目录）")
    parser.add_argument("--industry", default=None,
                        help="强制指定行业（如：半导体/芯片、医疗器械、生物医药等）")
    parser.add_argument("--mock-data", action="store_true",
                        help="使用模拟数据（跳过MCP调用，用于测试）")
    parser.add_argument("--input-json", metavar="FILE",
                        help="从JSON文件加载专利数据（格式：{专利号: markdown文本}）")

    args = parser.parse_args()

    # 收集专利号
    patent_numbers = []
    if args.patents:
        patent_numbers.extend(args.patents)
    if args.patents_file:
        with open(args.patents_file, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#"):
                    patent_numbers.extend([p.strip() for p in line.replace(",", " ").split() if p.strip()])

    if not patent_numbers:
        parser.print_help()
        sys.exit(0)

    patent_numbers = list(dict.fromkeys(patent_numbers))  # 去重保序
    print(f"[INFO] 待评审专利数量：{len(patent_numbers)}")

    # 加载数据
    patent_markdown_map = {}

    if args.input_json:
        with open(args.input_json, "r", encoding="utf-8") as f:
            patent_markdown_map = json.load(f)
        print(f"[INFO] 从JSON文件加载了 {len(patent_markdown_map)} 件专利数据")

    elif args.mock_data:
        # 模拟数据（测试用）
        for pn in patent_numbers:
            patent_markdown_map[pn] = f"""
# 专利详细信息

**Title:** 一种基于深度学习的图像分类方法及系统
**Assignee:** 测试科技股份有限公司
**Application Date:** 2021-03-15
**IPC:** H04N 19/00; G06T 7/00; G06F 18/214
**Legal Status:** 有效 / Active
**Family Count:** 5
**Jurisdictions:** CN US EP
**Cited By:** 12

## 摘要
本发明公开了一种基于深度学习的图像分类方法，包括：
数据预处理步骤、特征提取步骤、分类决策步骤。

## 权利要求
1. 一种图像分类方法，其特征在于，包括：
数据采集、特征提取、模型推理、结果输出。
2. 根据权利要求1所述的方法，其特征在于，所述特征提取采用卷积神经网络。
3. 一种计算机可读存储介质，存储有权利要求1所述方法的程序。
"""
        print(f"[INFO] 已生成 {len(patent_numbers)} 件模拟数据（--mock-data模式）")

    else:
        # 正常模式：提示Eureka AI通过MCP填充数据
        print("[INFO] 正常模式：请由Eureka通过智慧芽MCP获取数据后传入 patent_markdown_map")
        print("[INFO] 或使用 --mock-data 进行本地测试，--input-json 加载已缓存数据")
        for pn in patent_numbers:
            patent_markdown_map[pn] = ""

    # 执行评审
    output_path = run_grading(
        patent_numbers=patent_numbers,
        patent_markdown_map=patent_markdown_map,
        output_format=args.output_format,
        output_dir=args.output_dir,
        industry_override=args.industry,
    )

    print(f"[OK] 评审完成，输出文件：{output_path}")
    return output_path


if __name__ == "__main__":
    main()
